from __future__ import annotations
import argparse
import json
import os
import subprocess
import sys
import tempfile
import time
import logging
import hashlib
import pickle
import re
from concurrent.futures import ThreadPoolExecutor, as_completed
from functools import lru_cache, wraps
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Union, Callable
from datetime import datetime, timedelta
from dataclasses import dataclass, asdict
import dotenv
import requests
import torch
from icalendar import Calendar, Event
from faster_whisper import WhisperModel
from pyannote.audio import Pipeline
from tqdm import tqdm
import pandas as pd
import nltk
nltk.download('punkt_tab')
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.cluster import KMeans
# Optional heavy imports (avoided unless used)
try:
    import openai  # noqa: F401
except ImportError:
    openai = None  # type: ignore
try:
    from docx import Document
except ImportError:
    Document = None  # type: ignore
try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
except ImportError:
    canvas = None  # type: ignore
# ─────────────────────────────  CONFIG  ────────────────────────────── #
dotenv.load_dotenv()
# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("meeting_pipeline.log"),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)
# Configuration directory
CONFIG_DIR = Path(os.getenv("CONFIG_DIR", "~/.config/meeting_pipeline")).expanduser()
CONFIG_DIR.mkdir(parents=True, exist_ok=True)
CONFIG_FILE = CONFIG_DIR / "config.json"
CACHE_DIR = CONFIG_DIR / "cache"
CACHE_DIR.mkdir(exist_ok=True)
# Default configuration
DEFAULT_CONFIG = {
    "whisper_model": "base",
    "llm_backend": "auto",
    "openai_model": "gpt-4o-mini",
    "ollama_model": "llama3",
    "ollama_url": "http://localhost:11434",
    "num_speakers": None,
    "language": "en",
    "max_retries": 3,
    "retry_delay": 5,
    "parallel_workers": 4,
    "cache_enabled": True,
    "speaker_names": {},
    "default_due_days": 7,
    "export_formats": ["json", "txt"]
}
# Load or create configuration
def load_config() -> Dict[str, Any]:
    """Load configuration from file or create default."""
    if CONFIG_FILE.exists():
        try:
            with open(CONFIG_FILE, "r", encoding="utf-8") as f:
                config = json.load(f)
            # Merge with defaults to ensure all keys exist
            merged_config = DEFAULT_CONFIG.copy()
            merged_config.update(config)
            return merged_config
        except Exception as e:
            logger.warning(f"Failed to load config: {e}. Using defaults.")
    return DEFAULT_CONFIG.copy()
def save_config(config: Dict[str, Any]) -> None:
    """Save configuration to file."""
    try:
        with open(CONFIG_FILE, "w", encoding="utf-8") as f:
            json.dump(config, f, indent=2, ensure_ascii=False)
    except Exception as e:
        logger.error(f"Failed to save config: {e}")
# Initialize configuration
config = load_config()
# Device configuration
DEVICE = torch.device("cuda" if torch.cuda.is_available() else "cpu")
MIN_SLICE_SEC = 0.20  # avoid zero‑length clips
torch.backends.cuda.matmul.allow_tf32 = True
torch.backends.cudnn.allow_tf32 = True
# Download required NLTK data
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt', quiet=True)
try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('stopwords', quiet=True)
# Supported file formats
AUDIO_FORMATS = ['.mp3', '.wav', '.flac', '.m4a', '.aac', '.ogg', '.wma']
VIDEO_FORMATS = ['.mp4', '.mov', '.avi', '.mkv', '.webm', '.flv', '.wmv']
SUPPORTED_FORMATS = AUDIO_FORMATS + VIDEO_FORMATS
# ────────────────────────────  HELPERS  ───────────────────────────── #
def _assert_file(p: str | Path) -> Path:
    path = Path(p)
    if not path.exists():
        raise FileNotFoundError(path)
    return path.resolve()
def _run_ffmpeg(cmd: list[str]):
    try:
        subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    except FileNotFoundError as exc:
        raise RuntimeError("ffmpeg executable not found on PATH") from exc
    except subprocess.CalledProcessError as exc:
        sys.stderr.write("FFmpeg error:\n" + exc.stderr.decode() + "\n")
        raise
def retry(max_retries: int = 3, delay: float = 5.0, exceptions: Tuple = (Exception,)):
    """Decorator for retrying a function call."""
    def decorator(func):
        @wraps(func)
        def wrapper(*args, **kwargs):
            last_exception = None
            for attempt in range(1, max_retries + 1):
                try:
                    return func(*args, **kwargs)
                except exceptions as e:
                    last_exception = e
                    if attempt < max_retries:
                        logger.warning(f"Attempt {attempt}/{max_retries} failed: {e}. Retrying in {delay} seconds...")
                        time.sleep(delay)
                    else:
                        logger.error(f"All {max_retries} attempts failed.")
            raise last_exception
        return wrapper
    return decorator
def get_file_hash(file_path: str | Path) -> str:
    """Calculate SHA256 hash of a file."""
    hash_sha256 = hashlib.sha256()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(4096), b""):
            hash_sha256.update(chunk)
    return hash_sha256.hexdigest()
def cache_result(cache_dir: Path = CACHE_DIR):
    """Decorator for caching function results."""
    def decorator(func):
        @wraps(func)
        def wrapper(*args, **kwargs):
            if not config.get("cache_enabled", True):
                return func(*args, **kwargs)
                
            # Create a unique cache key based on function name and arguments
            args_str = "_".join(str(arg) for arg in args)
            kwargs_str = "_".join(f"{k}={v}" for k, v in sorted(kwargs.items()))
            cache_key = f"{func.__name__}_{hashlib.md5((args_str + kwargs_str).encode()).hexdigest()}"
            cache_file = cache_dir / f"{cache_key}.pkl"
            
            # Check if result is cached
            if cache_file.exists():
                try:
                    with open(cache_file, "rb") as f:
                        result = pickle.load(f)
                    logger.info(f"Using cached result for {func.__name__}")
                    return result
                except Exception as e:
                    logger.warning(f"Failed to load cached result: {e}")
            
            # Compute and cache the result
            result = func(*args, **kwargs)
            try:
                with open(cache_file, "wb") as f:
                    pickle.dump(result, f)
                logger.info(f"Cached result for {func.__name__}")
            except Exception as e:
                logger.warning(f"Failed to cache result: {e}")
                
            return result
        return wrapper
    return decorator
def is_supported_format(file_path: str | Path) -> bool:
    """Check if file format is supported."""
    file_ext = Path(file_path).suffix.lower()
    return file_ext in SUPPORTED_FORMATS
def is_video_format(file_path: str | Path) -> bool:
    """Check if file is a video format."""
    file_ext = Path(file_path).suffix.lower()
    return file_ext in VIDEO_FORMATS
def convert_to_audio(input_path: str | Path, output_path: str | Path = None) -> Path:
    """Convert video or audio file to WAV format using ffmpeg."""
    input_path = Path(input_path)
    
    if output_path is None:
        output_path = input_path.with_suffix('.wav')
    else:
        output_path = Path(output_path)
    
    # Use ffmpeg to convert to WAV
    cmd = [
        "ffmpeg", "-y", "-i", str(input_path),
        "-ac", "1", "-ar", "16000", "-loglevel", "error", str(output_path),
    ]
    _run_ffmpeg(cmd)
    
    return output_path
# ───────────────────────  FILE METADATA EXTRACTION  ─────────────────── #
@dataclass
class MeetingMetadata:
    file_path: str
    file_size: int
    created_time: float
    modified_time: float
    duration: Optional[float] = None
    tags: Optional[Dict[str, str]] = None
    meeting_date: Optional[datetime] = None
    meeting_topic: Optional[str] = None
    detected_language: Optional[str] = None
    file_type: Optional[str] = None  # 'audio' or 'video'
def get_file_metadata(file_path: str | Path, meeting_date=None) -> MeetingMetadata:
    """Extract metadata from the file including creation and modification dates."""
    path = _assert_file(file_path)
    stat = path.stat()
    
    # Determine file type
    file_type = "video" if is_video_format(path) else "audio"
    
    metadata = MeetingMetadata(
        file_path=str(path),
        file_size=stat.st_size,
        created_time=stat.st_ctime,  # Creation time (platform-dependent)
        modified_time=stat.st_mtime,  # Last modification time
        file_type=file_type
    )
    
    # Try to extract more metadata using ffmpeg if available
    try:
        cmd = ["ffprobe", "-v", "quiet", "-print_format", "json", "-show_format", str(path)]
        result = subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
        ffprobe_data = json.loads(result.stdout)
        
        if "format" in ffprobe_data:
            format_info = ffprobe_data["format"]
            if "tags" in format_info:
                metadata.tags = format_info["tags"]
            
            # Add duration if available
            if "duration" in format_info:
                metadata.duration = float(format_info["duration"])
    except (subprocess.CalledProcessError, FileNotFoundError, json.JSONDecodeError):
        # ffprobe not available or failed, just use basic file stat
        pass
    
    # use file creation time
    if meeting_date:
        metadata.meeting_date = meeting_date
    else:
        metadata.meeting_date = datetime.fromtimestamp(metadata.created_time)
    print(f"Meeting is detected as in {metadata.meeting_date}")
    return metadata
# ──────────────────────────  LOCAL WHISPER  ────────────────────────── #
@lru_cache(maxsize=2)
def _load_whisper(model_size: str = "base") -> WhisperModel:
    print(f"Loading faster‑whisper '{model_size}' on {DEVICE.type} …")
    return WhisperModel(model_size, device=DEVICE.type, compute_type="int8")
def _transcribe_local(audio_slice: Path, whisper_model: WhisperModel, language: str = None) -> Tuple[str, str]:
    """Transcribe audio slice and return text and detected language."""
    if language:
        segments, _ = whisper_model.transcribe(str(audio_slice), vad_filter=False, language=language)
    else:
        segments, _ = whisper_model.transcribe(str(audio_slice), vad_filter=False)
    
    text = " ".join(seg.text for seg in segments).strip()
    
    # Get language from the first segment if not provided
    if language is None:
        try:
            first_segment = next(segments)
            detected_lang = first_segment.language
        except StopIteration:
            detected_lang = "en"  # Default to English if no segments
    else:
        detected_lang = language
    
    return text, detected_lang
# ────────────────────────  SPEAKER DIARIZATION  ────────────────────── #
@cache_result()
@retry(max_retries=config.get("max_retries", 3), delay=config.get("retry_delay", 5))
def diarize_audio(audio_path: str | Path, num_speakers: Optional[int] = None) -> List[Dict[str, Any]]:
    print("Running speaker diarization …")
    try:
        pipeline = Pipeline.from_pretrained(
            "pyannote/speaker-diarization-3.1",
            use_auth_token=os.getenv("HUGGINGFACE_TOKEN"),
        )
        pipeline.to(DEVICE)
        diarization = pipeline(_assert_file(audio_path), num_speakers=num_speakers)
        segs: List[Dict[str, Any]] = []
        for turn, _, spk in diarization.itertracks(yield_label=True):
            segs.append({"start": turn.start, "end": turn.end, "speaker": spk})
        print(f"Diarization produced {len(segs)} segments")
        return segs
    except Exception as e:
        logger.error(f"Diarization failed: {e}")
        # Fallback to single speaker if diarization fails
        logger.warning("Falling back to single speaker diarization")
        return [{"start": 0, "end": get_file_metadata(audio_path).duration or 0, "speaker": "SPEAKER_00"}]
# ─────────────────────────  TRANSCRIPTION  ─────────────────────────── #
def _extract_segment(src: Path, start: float, end: float) -> Path:
    duration = max(end - start, MIN_SLICE_SEC)
    tmp = tempfile.NamedTemporaryFile(suffix=".wav", delete=False)
    cmd = [
        "ffmpeg", "-y", "-i", str(src),
        "-ss", f"{start}", "-t", f"{duration}",
        "-ac", "1", "-ar", "16000", "-loglevel", "error", tmp.name,
    ]
    _run_ffmpeg(cmd)
    return Path(tmp.name)
def _transcribe_segment(args: Tuple[Path, Path, WhisperModel, str]) -> Dict[str, Any]:
    """Helper function for parallel transcription of a segment."""
    src, seg, whisper, language = args
    if seg["end"] - seg["start"] < 0.05:
        return {"speaker": seg["speaker"], "text": "", "language": language}
    
    clip = _extract_segment(src, seg["start"], seg["end"])
    try:
        text, detected_lang = _transcribe_local(clip, whisper, language)
        return {
            "speaker": seg["speaker"],
            "text": text,
            "start": seg["start"],
            "end": seg["end"],
            "language": detected_lang
        }
    finally:
        clip.unlink(missing_ok=True)
@cache_result()
def compile_transcript(
    audio_path: str | Path,
    segments: List[Dict[str, Any]],
    whisper_size: str = "base",
    language: str = "en"  # default language
) -> Tuple[str, Dict[str, str], str]:
    """Compile transcript with parallel processing and language detection."""
    whisper = _load_whisper(whisper_size)
    src = _assert_file(audio_path)
    
    # Prepare arguments for parallel processing
    args_list = [(src, seg, whisper, language) for seg in segments]
    
    # Process segments in parallel
    results = []
    with ThreadPoolExecutor(max_workers=config.get("parallel_workers", 4)) as executor:
        futures = [executor.submit(_transcribe_segment, args) for args in args_list]
        
        for future in tqdm(as_completed(futures), total=len(futures), desc="Transcribing", unit="seg"):
            try:
                result = future.result()
                if result["text"]:
                    results.append(result)
            except Exception as e:
                logger.error(f"Error transcribing segment: {e}")
    
    # Build transcript and collect language information
    transcript_lines = []
    speaker_languages = {}
    dominant_language = None
    language_counts = {}
    
    for result in results:
        speaker = result["speaker"]
        text = result["text"]
        lang = result["language"]
        
        transcript_lines.append(f"{speaker}: {text}")
        
        # Track languages per speaker
        if speaker not in speaker_languages:
            speaker_languages[speaker] = {}
        speaker_languages[speaker][lang] = speaker_languages[speaker].get(lang, 0) + 1
        
        # Track overall language counts
        language_counts[lang] = language_counts.get(lang, 0) + 1
    
    # Determine dominant language
    if language_counts:
        dominant_language = max(language_counts, key=language_counts.get)
    
    # For each speaker, determine their dominant language
    for speaker, langs in speaker_languages.items():
        speaker_languages[speaker] = max(langs, key=langs.get)
    
    transcript = "\n".join(transcript_lines)
    return transcript, speaker_languages, dominant_language
# ───────────────────  TEXT ANALYSIS UTILITIES  ────────────────────── #
def extract_keywords(text: str, max_keywords: int = 10) -> List[str]:
    """Extract keywords from text using TF-IDF."""
    try:
        # Tokenize and remove stopwords
        stop_words = set(stopwords.words('english'))
        tokens = word_tokenize(text.lower())
        filtered_tokens = [word for word in tokens if word.isalpha() and word not in stop_words and len(word) > 2]
        
        # Join tokens back into text
        filtered_text = ' '.join(filtered_tokens)
        
        # Use TF-IDF to extract keywords
        vectorizer = TfidfVectorizer(max_features=max_keywords, ngram_range=(1, 2))
        tfidf_matrix = vectorizer.fit_transform([filtered_text])
        feature_names = vectorizer.get_feature_names_out()
        scores = tfidf_matrix.toarray()[0]
        
        # Sort by score and return top keywords
        keywords = [feature_names[i] for i in scores.argsort()[-max_keywords:][::-1]]
        return keywords
    except Exception as e:
        logger.error(f"Error extracting keywords: {e}")
        return []
def detect_topic(transcript: str, metadata: MeetingMetadata) -> str:
    """Detect meeting topic using simple heuristics and filename."""
    # Try to extract topic from filename
    filename = Path(metadata.file_path).stem.lower()
    
    # Common meeting topic keywords
    topic_keywords = {
        "planning": ["plan", "planning", "strategy", "roadmap"],
        "review": ["review", "retrospective", "feedback", "evaluation"],
        "project": ["project", "sprint", "milestone", "deliverable"],
        "budget": ["budget", "finance", "cost", "financial"],
        "hiring": ["hiring", "interview", "recruitment", "candidate"],
        "marketing": ["marketing", "campaign", "promotion", "brand"],
        "sales": ["sales", "revenue", "customer", "deal"],
        "technical": ["technical", "development", "engineering", "architecture"],
        "product": ["product", "feature", "user", "experience"],
    }
    
    # Check filename for topic keywords
    for topic, keywords in topic_keywords.items():
        if any(keyword in filename for keyword in keywords):
            return topic.capitalize()
    
    # Check transcript for topic keywords
    for topic, keywords in topic_keywords.items():
        if any(keyword in transcript.lower() for keyword in keywords):
            return topic.capitalize()
    
    # Default topic
    return "General"
def identify_speakers(transcript: str, speaker_names: Dict[str, str]) -> Dict[str, str]:
    """Identify speakers based on transcript content and provided names."""
    # Simple heuristic: look for self-introductions like "I'm [Name]" or "This is [Name]"
    speaker_patterns = {}
    
    for line in transcript.split('\n'):
        if not line.strip():
            continue
            
        speaker, text = line.split(':', 1)
        text = text.strip()
        
        # Look for self-introduction patterns
        patterns = [
            r"I'm ([A-Z][a-z]+)",
            r"I am ([A-Z][a-z]+)",
            r"This is ([A-Z][a-z]+)",
            r"([A-Z][a-z]+) here",
            r"([A-Z][a-z]+) speaking"
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text)
            if match:
                name = match.group(1)
                speaker_patterns[speaker] = name
                break
    
    # Merge with provided speaker names
    identified_speakers = {**speaker_names, **speaker_patterns}
    return identified_speakers
# ───────────────────  LLM MEETING INTELLIGENCE  ────────────────────── #
SYSTEM_PROMPT = (
    "You are a senior executive assistant. Given a verbatim, speaker‑labelled transcript of a meeting, "
    "which includes the meeting date at the beginning, respond in valid JSON with keys: summary (3‑5 bullets), "
    "decisions, action_items (array of {task, owner, due_date}), open_questions. "
    "For due dates, use specific dates (YYYY-MM-DD) when possible, or relative references like 'tomorrow', "
    "'next week', 'end of month' if the exact date isn't mentioned but can be inferred from context. "
    "Also include a 'meeting_topic' field with a brief description of the meeting's main topic. "
    "Return only the JSON."
)
@retry(max_retries=config.get("max_retries", 3), delay=config.get("retry_delay", 5))
def _call_openai(transcript: str, model: str) -> Dict[str, Any]:
    if openai is None:
        raise RuntimeError("openai package not installed")
    if not os.getenv("OPENAI_API_KEY"):
        raise RuntimeError("OPENAI_API_KEY not set")
    resp = openai.chat.completions.create(
        model=model,
        response_format={"type": "json_object"},
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": transcript},
        ],
    )
    return json.loads(resp.choices[0].message.content)
@retry(max_retries=config.get("max_retries", 3), delay=config.get("retry_delay", 5))
def _call_ollama(transcript: str, model: str = "llama3", ollama_url: str = None) -> Dict[str, Any]:
    if ollama_url is None:
        ollama_url = config.get("ollama_url", "http://localhost:11434")
    
    url = f"{ollama_url}/api/chat"
    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": transcript},
        ],
        "format": "json",
        "stream": False,
        "options": {
            "num_ctx": 32768
        }
    }
    
    # First check if Ollama is running
    try:
        r = requests.get(f"{ollama_url}/api/tags", timeout=5)
        r.raise_for_status()
    except requests.exceptions.RequestException as e:
        logger.error(f"Ollama is not running or not accessible at {ollama_url}: {e}")
        raise RuntimeError(f"Ollama is not running or not accessible at {ollama_url}. Please start Ollama or check the URL in your configuration.")
    
    # Then try to make the chat request
    try:
        r = requests.post(url, json=payload, timeout=300)
        r.raise_for_status()
        content = r.json()["message"]
        return json.loads(content["content"]) if isinstance(content, dict) else json.loads(content)
    except requests.exceptions.RequestException as e:
        logger.error(f"Ollama request failed: {e}")
        if e.response and e.response.status_code == 404:
            raise RuntimeError(f"Ollama API endpoint not found at {url}. Please check your Ollama installation and version.")
        raise RuntimeError(f"Failed to connect to Ollama: {e}")
    except json.JSONDecodeError as e:
        logger.error(f"Failed to parse Ollama response: {e}")
        raise RuntimeError(f"Failed to parse Ollama response: {e}")
def analyse_meeting(
    transcript: str,
    backend: str = "auto",
    openai_model: str = "gpt-4o-mini",
    ollama_model: str = "llama3",
    ollama_url: str = None
) -> Dict[str, Any]:
    last_exception = None
    
    if backend == "openai":
        try:
            return _call_openai(transcript, openai_model)
        except Exception as e:
            logger.error(f"OpenAI analysis failed: {e}")
            last_exception = e
    
    if backend == "ollama":
        try:
            return _call_ollama(transcript, ollama_model, ollama_url)
        except Exception as e:
            logger.error(f"Ollama analysis failed: {e}")
            last_exception = e
    
    # auto mode
    if os.getenv("OPENAI_API_KEY") and openai is not None:
        try:
            return _call_openai(transcript, openai_model)
        except Exception as e:
            logger.error(f"OpenAI fallback failed: {e}")
            last_exception = e
    
    try:
        return _call_ollama(transcript, ollama_model, ollama_url)
    except Exception as e:
        logger.error(f"Ollama fallback failed: {e}")
        last_exception = e
    
    # If all attempts failed, return a basic structure
    logger.error("All LLM backends failed. Returning basic analysis.")
    return {
        "summary": ["Meeting analysis failed due to technical issues."],
        "decisions": [],
        "action_items": [],
        "open_questions": [],
        "meeting_topic": "Unknown"
    }
# ─────────────────────────────  ICALENDAR  ──────────────────────────── #
def parse_relative_date(date_str: str, meeting_date: datetime) -> datetime:
    """Parse relative date expressions like 'tomorrow', 'next week' etc."""
    date_str = date_str.lower()
    
    # Handle common relative expressions
    if "tomorrow" in date_str:
        return meeting_date + timedelta(days=1)
    elif "next week" in date_str or "in a week" in date_str:
        return meeting_date + timedelta(weeks=1)
    elif "in two weeks" in date_str or "2 weeks" in date_str:
        return meeting_date + timedelta(weeks=2)
    elif "end of month" in date_str:
        if meeting_date.month == 12:
            return meeting_date.replace(year=meeting_date.year + 1, month=1, day=1) - timedelta(days=1)
        else:
            return meeting_date.replace(month=meeting_date.month + 1, day=1) - timedelta(days=1)
    elif "next month" in date_str:
        if meeting_date.month == 12:
            return meeting_date.replace(year=meeting_date.year + 1, month=1, day=1)
        else:
            return meeting_date.replace(month=meeting_date.month + 1, day=1)
    elif "in a month" in date_str:
        return meeting_date + timedelta(days=30)
    
    # If no pattern matches, return meeting date + default days
    return meeting_date + timedelta(days=config.get("default_due_days", 7))
def build_ics(action_items: List[Dict[str, Any]], filename: str = "meeting_tasks.ics", meeting_date=None) -> Optional[Path]:
    if not action_items:
        logger.warning("No action items to add to calendar")
        return None
        
    cal = Calendar()
    cal.add("prodid", "-//Meeting‑Pipeline//")
    cal.add("version", "2.0")
    added = False
    
    for item in action_items:
        due = item.get("due_date")
        
        # If no due date but we have meeting date, set a default due date
        if not due and meeting_date:
            default_due = meeting_date + timedelta(days=config.get("default_due_days", 7))
            due = default_due.strftime("%Y-%m-%d")
            item["due_date"] = due  # Update the item with the inferred due date
        
        if not due:
            continue
            
        # Parse due date if it's a string
        try:
            if isinstance(due, str):
                # Handle various date formats
                for fmt in ("%Y-%m-%d", "%m/%d/%Y", "%d.%m.%Y", "%B %d, %Y"):
                    try:
                        due_date = datetime.strptime(due, fmt)
                        break
                    except ValueError:
                        continue
                else:
                    # If no format matched, try to extract relative dates
                    due_date = parse_relative_date(due, meeting_date or datetime.now())
            else:
                due_date = due
        except (ValueError, TypeError, AttributeError):
            continue
            
        evt = Event()
        evt.add("summary", item.get("task", "Action Item"))
        evt.add("dtstart", due_date)
        evt.add("dtend", due_date)
        if owner := item.get("owner"):
            evt.add("organizer", owner)
        cal.add_component(evt)
        added = True
    
    if added:
        path = Path(filename)
        path.write_bytes(cal.to_ical())
        return path
    return None
# ───────────────────────  EXPORT UTILITIES  ─────────────────────────── #
def export_to_txt(data: Dict[str, Any], filename: str) -> Path:
    """Export meeting results to a text file."""
    path = Path(filename)
    with open(path, "w", encoding="utf-8") as f:
        f.write("MEETING SUMMARY\n")
        f.write("=" * 50 + "\n\n")
        
        if "metadata" in data and "meeting_date" in data["metadata"]:
            f.write(f"Date: {data['metadata']['meeting_date'].strftime('%Y-%m-%d %H:%M')}\n")
        
        if "meeting_topic" in data:
            f.write(f"Topic: {data['meeting_topic']}\n")
        
        f.write("\nSUMMARY\n")
        f.write("-" * 20 + "\n")
        if "summary" in data:
            for point in data["summary"]:
                f.write(f"• {point}\n")
        
        f.write("\nDECISIONS\n")
        f.write("-" * 20 + "\n")
        if "decisions" in data:
            for decision in data["decisions"]:
                f.write(f"• {decision}\n")
        
        f.write("\nACTION ITEMS\n")
        f.write("-" * 20 + "\n")
        if "action_items" in data:
            for item in data["action_items"]:
                f.write(f"• {item.get('task', 'Unknown task')}\n")
                if "owner" in item:
                    f.write(f"  Owner: {item['owner']}\n")
                if "due_date" in item:
                    f.write(f"  Due: {item['due_date']}\n")
                f.write("\n")
        
        f.write("\nOPEN QUESTIONS\n")
        f.write("-" * 20 + "\n")
        if "open_questions" in data:
            for question in data["open_questions"]:
                f.write(f"• {question}\n")
        
        if "keywords" in data:
            f.write("\nKEYWORDS\n")
            f.write("-" * 20 + "\n")
            f.write(", ".join(data["keywords"]))
    
    return path
def export_to_docx(data: Dict[str, Any], filename: str) -> Optional[Path]:
    """Export meeting results to a DOCX file."""
    if Document is None:
        logger.warning("python-docx not installed. Cannot export to DOCX.")
        return None
        
    path = Path(filename)
    document = Document()
    
    # Title
    document.add_heading('Meeting Summary', 0)
    
    # Metadata
    if "metadata" in data and "meeting_date" in data["metadata"]:
        document.add_paragraph(f"Date: {data['metadata']['meeting_date'].strftime('%Y-%m-%d %H:%M')}")
    
    if "meeting_topic" in data:
        document.add_paragraph(f"Topic: {data['meeting_topic']}")
    
    # Summary
    document.add_heading('Summary', level=1)
    if "summary" in data:
        for point in data["summary"]:
            document.add_paragraph(f"• {point}")
    
    # Decisions
    document.add_heading('Decisions', level=1)
    if "decisions" in data:
        for decision in data["decisions"]:
            document.add_paragraph(f"• {decision}")
    
    # Action Items
    document.add_heading('Action Items', level=1)
    if "action_items" in data:
        for item in data["action_items"]:
            p = document.add_paragraph(f"• {item.get('task', 'Unknown task')}")
            if "owner" in item:
                p.add_run(f"\nOwner: {item['owner']}")
            if "due_date" in item:
                p.add_run(f"\nDue: {item['due_date']}")
            document.add_paragraph()
    
    # Open Questions
    document.add_heading('Open Questions', level=1)
    if "open_questions" in data:
        for question in data["open_questions"]:
            document.add_paragraph(f"• {question}")
    
    # Keywords
    if "keywords" in data:
        document.add_heading('Keywords', level=1)
        document.add_paragraph(", ".join(data["keywords"]))
    
    document.save(path)
    return path
def export_to_pdf(data: Dict[str, Any], filename: str) -> Optional[Path]:
    """Export meeting results to a PDF file."""
    if canvas is None:
        logger.warning("reportlab not installed. Cannot export to PDF.")
        return None
        
    path = Path(filename)
    doc = SimpleDocTemplate(str(path), pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    # Title
    story.append(Paragraph("Meeting Summary", styles["Title"]))
    
    # Metadata
    if "metadata" in data and "meeting_date" in data["metadata"]:
        story.append(Paragraph(f"Date: {data['metadata']['meeting_date'].strftime('%Y-%m-%d %H:%M')}", styles["Normal"]))
    
    if "meeting_topic" in data:
        story.append(Paragraph(f"Topic: {data['meeting_topic']}", styles["Normal"]))
    
    story.append(Spacer(1, 12))
    
    # Summary
    story.append(Paragraph("Summary", styles["Heading1"]))
    if "summary" in data:
        for point in data["summary"]:
            story.append(Paragraph(f"• {point}", styles["Normal"]))
    
    story.append(Spacer(1, 12))
    
    # Decisions
    story.append(Paragraph("Decisions", styles["Heading1"]))
    if "decisions" in data:
        for decision in data["decisions"]:
            story.append(Paragraph(f"• {decision}", styles["Normal"]))
    
    story.append(Spacer(1, 12))
    
    # Action Items
    story.append(Paragraph("Action Items", styles["Heading1"]))
    if "action_items" in data:
        for item in data["action_items"]:
            story.append(Paragraph(f"• {item.get('task', 'Unknown task')}", styles["Normal"]))
            if "owner" in item:
                story.append(Paragraph(f"Owner: {item['owner']}", styles["Normal"]))
            if "due_date" in item:
                story.append(Paragraph(f"Due: {item['due_date']}", styles["Normal"]))
            story.append(Spacer(1, 6))
    
    story.append(Spacer(1, 12))
    
    # Open Questions
    story.append(Paragraph("Open Questions", styles["Heading1"]))
    if "open_questions" in data:
        for question in data["open_questions"]:
            story.append(Paragraph(f"• {question}", styles["Normal"]))
    
    # Keywords
    if "keywords" in data:
        story.append(Spacer(1, 12))
        story.append(Paragraph("Keywords", styles["Heading1"]))
        story.append(Paragraph(", ".join(data["keywords"]), styles["Normal"]))
    
    doc.build(story)
    return path
    
# ───────────────────────────  GRADIO UI  ───────────────────────────── #
def _launch_ui():
    import gradio as gr  # local import to avoid mandatory dependency
    
    def process(
        audio_file, meeting_date:float, whisper_size, num_speakers, llm_backend, openai_model, ollama_model, ollama_url, language, export_formats, progress=gr.Progress()
    ):
        meeting_date = datetime.fromtimestamp(meeting_date) if meeting_date else None
        
        if audio_file is None:
            return "Please upload a file.", None, None, None, None, None
        
        # Check if file format is supported
        if not is_supported_format(audio_file):
            return f"Unsupported file format. Supported formats: {', '.join(SUPPORTED_FORMATS)}", None, None, None, None, None
        
        with tempfile.NamedTemporaryFile(delete=False, suffix=Path(audio_file).suffix) as tmp:
            tmp.write(Path(audio_file).read_bytes())
            tmp_path = tmp.name
        
        # Create a progress tracker
        progress_steps = [
            ("Extracting metadata", 0.1),
            ("Converting video to audio" if is_video_format(tmp_path) else "Preparing audio", 0.2),
            ("Running speaker diarization", 0.4),
            ("Transcribing audio", 0.6),
            ("Analyzing meeting content", 0.8),
            ("Generating exports", 0.95),
            ("Complete", 1.0)
        ]
        
        progress(0, desc="Starting...")
        
        # Step 1: Extract metadata
        progress(0.1, desc=progress_steps[0][0])
        metadata = get_file_metadata(tmp_path, meeting_date)
        
        # Step 2: Convert video to audio if needed
        if is_video_format(tmp_path):
            progress(0.2, desc=progress_steps[1][0])
            with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as tmp_audio:
                tmp_audio_path = tmp_audio.name
            convert_to_audio(tmp_path, tmp_audio_path)
            audio_to_process = tmp_audio_path
        else:
            audio_to_process = tmp_path
        
        # Step 3: Diarize audio
        progress(0.4, desc=progress_steps[2][0])
        segments = diarize_audio(audio_to_process, num_speakers=num_speakers if num_speakers else None)
        
        # Step 4: Transcribe audio
        progress(0.6, desc=progress_steps[3][0])
        transcript, speaker_languages, dominant_language = compile_transcript(
            audio_to_process, segments, whisper_size, language=language
        )
        
        # Update metadata with detected language
        metadata.detected_language = dominant_language
        
        # Detect meeting topic
        meeting_topic = detect_topic(transcript, metadata)
        metadata.meeting_topic = meeting_topic
        
        # Identify speakers
        speaker_names = identify_speakers(transcript, config.get("speaker_names", {}))
        
        # Add meeting date context to the transcript for LLM analysis
        meeting_date_str = ""
        meeting_date = metadata.meeting_date
        if meeting_date:
            meeting_date_str = f"Meeting date: {meeting_date.strftime('%Y-%m-%d %H:%M')}"
        
        if meeting_date_str:
            transcript_with_date = f"{meeting_date_str}\n\nMeeting topic: {meeting_topic}\n\n{transcript}"
        else:
            transcript_with_date = f"Meeting topic: {meeting_topic}\n\n{transcript}"
        
        # Step 5: Analyze meeting with LLM
        progress(0.8, desc=progress_steps[4][0])
        insights = analyse_meeting(
            transcript_with_date, 
            backend=llm_backend, 
            openai_model=openai_model, 
            ollama_model=ollama_model,
            ollama_url=ollama_url
        )
        
        # Extract keywords
        keywords = extract_keywords(transcript)
        
        # Compile final results
        results = {
            "metadata": asdict(metadata),
            "meeting_topic": meeting_topic,
            "speaker_languages": speaker_languages,
            "speaker_names": speaker_names,
            "keywords": keywords,
            "raw_transcript": transcript,
            **insights
        }
        
        # Generate calendar
        if (ics := build_ics(insights.get("action_items", []), meeting_date=meeting_date)):
            results["ics_file"] = str(ics)
        
        # Step 6: Export to requested formats
        progress(0.95, desc=progress_steps[5][0])
        base_path = Path(audio_to_process).stem
        for fmt in export_formats:
            if fmt.lower() == "json":
                json_path = Path(f"{base_path}_meeting_summary.json")
                with open(json_path, "w", encoding="utf-8") as f:
                    json.dump(results, f, indent=2, ensure_ascii=False, default=str)
                results["json_file"] = str(json_path)
            
            elif fmt.lower() == "txt":
                txt_path = Path(f"{base_path}_meeting_summary.txt")
                export_to_txt(results, str(txt_path))
                results["txt_file"] = str(txt_path)
            
            elif fmt.lower() == "docx":
                docx_path = Path(f"{base_path}_meeting_summary.docx")
                export_to_docx(results, str(docx_path))
                results["docx_file"] = str(docx_path)
            
            elif fmt.lower() == "pdf":
                pdf_path = Path(f"{base_path}_meeting_summary.pdf")
                export_to_pdf(results, str(pdf_path))
                results["pdf_file"] = str(pdf_path)
        
        # Clean up temporary files
        try:
            Path(tmp_path).unlink()
            if is_video_format(tmp_path) and 'tmp_audio_path' in locals():
                Path(tmp_audio_path).unlink()
        except Exception as e:
            logger.warning(f"Failed to remove temporary files: {e}")
        
        progress(1.0, desc=progress_steps[6][0])
        
        json_out = json.dumps(results, indent=2, ensure_ascii=False, default=str)
        ics_file = results.get("ics_file")
        json_file = results.get("json_file")
        txt_file = results.get("txt_file")
        docx_file = results.get("docx_file")
        pdf_file = results.get("pdf_file")
        
        return json_out, json_file, txt_file, docx_file, pdf_file, ics_file
    
    def save_config_ui(
        whisper_model, llm_backend, openai_model, ollama_model, ollama_url, language, 
        max_retries, retry_delay, parallel_workers, cache_enabled, default_due_days
    ):
        new_config = config.copy()
        new_config.update({
            "whisper_model": whisper_model,
            "llm_backend": llm_backend,
            "openai_model": openai_model,
            "ollama_model": ollama_model,
            "ollama_url": ollama_url,
            "language": language,
            "max_retries": max_retries,
            "retry_delay": retry_delay,
            "parallel_workers": parallel_workers,
            "cache_enabled": cache_enabled,
            "default_due_days": default_due_days
        })
        save_config(new_config)
        return f"Configuration saved to {CONFIG_FILE}", json.dumps(new_config, indent=2)
    
    def get_file_info(file_path):
        if file_path is None:
            return "No file uploaded"
        
        try:
            metadata = get_file_metadata(file_path)
            file_type = "Video" if metadata.file_type == "video" else "Audio"
            file_size_mb = metadata.file_size / (1024 * 1024)
            
            info = f"**File Type:** {file_type}\n"
            info += f"**File Size:** {file_size_mb:.2f} MB\n"
            
            if metadata.duration:
                hours, remainder = divmod(metadata.duration, 3600)
                minutes, seconds = divmod(remainder, 60)
                duration_str = f"{int(hours)}h {int(minutes)}m {int(seconds)}s" if hours else f"{int(minutes)}m {int(seconds)}s"
                info += f"**Duration:** {duration_str}\n"
            
            if metadata.meeting_date:
                info += f"**Date:** {metadata.meeting_date.strftime('%Y-%m-%d %H:%M')}\n"
            
            return info
        except Exception as e:
            return "Error loading file"
    
    def _postprocess(data, json_file, txt_file, docx_file, pdf_file, ics_path):
        updates = {}
        
        # Parse JSON to extract summary, transcript, and action items
        try:
            # Format summary
            summary = "## Meeting Summary\n\n"
            if "meeting_topic" in data:
                summary += f"**Topic:** {data['meeting_topic']}\n\n"
            
            if "metadata" in data and "meeting_date" in data["metadata"]:
                meeting_date = data["metadata"]["meeting_date"]
                if isinstance(meeting_date, str):
                    try:
                        meeting_date = datetime.fromisoformat(meeting_date.replace('Z', '+00:00'))
                    except:
                        meeting_date = meeting_date
                if isinstance(meeting_date, datetime):
                    summary += f"**Date:** {meeting_date.strftime('%Y-%m-%d %H:%M')}\n\n"
            
            if "summary" in data:
                summary += "### Key Points\n\n"
                for point in data["summary"]:
                    summary += f"- {point}\n"
            
            # Format transcript
            transcript = "## Transcript\n\n"
            if "raw_transcript" in data:
                transcript += data["raw_transcript"]
            
            # Format action items
            actions = "## Action Items\n\n"
            if "action_items" in data and data["action_items"]:
                for item in data["action_items"]:
                    actions += f"### {item.get('task', 'Unknown task')}\n"
                    if "owner" in item:
                        actions += f"**Owner:** {item['owner']}\n"
                    if "due_date" in item:
                        actions += f"**Due Date:** {item['due_date']}\n"
                    actions += "\n"
            else:
                actions += "No action items identified.\n"
            
            updates[summary_out] = summary
            updates[transcript_out] = transcript
            updates[actions_out] = actions
        except Exception as e:
            logger.error(f"Error formatting output: {e}")
            updates[summary_out] = "Error processing results"
            updates[transcript_out] = "Error processing results"
            updates[actions_out] = "Error processing results"
        
        # Update file downloads
        if json_file and Path(json_file).exists():
            updates[json_download] = gr.update(value=json_file, visible=True)
        else:
            updates[json_download] = gr.update(visible=False)
        
        if txt_file and Path(txt_file).exists():
            updates[txt_download] = gr.update(value=txt_file, visible=True)
        else:
            updates[txt_download] = gr.update(visible=False)
        
        if docx_file and Path(docx_file).exists():
            updates[docx_download] = gr.update(value=docx_file, visible=True)
        else:
            updates[docx_download] = gr.update(visible=False)
        
        if pdf_file and Path(pdf_file).exists():
            updates[pdf_download] = gr.update(value=pdf_file, visible=True)
        else:
            updates[pdf_download] = gr.update(visible=False)
        
        if ics_path and Path(ics_path).exists():
            updates[ics_download] = gr.update(value=ics_path, visible=True)
        else:
            updates[ics_download] = gr.update(visible=False)
        
        updates[json_out] = data
        return updates
    
    # Custom CSS for better UI
    custom_css = """
    .main-header {
        text-align: center;
        margin-bottom: 20px;
    }
    .file-info {
        background-color: darkgray;
        border-radius: 8px;
        padding: 10px;
        margin-bottom: 15px;
    }
    .format-hint {
        font-size: 0.9em;
        color: #666;
        margin-top: 5px;
    }
    .progress-container {
        margin: 15px 0;
    }
    """
    
    with gr.Blocks(title="Meeting Pipeline", css=custom_css) as demo:
        gr.Markdown("# 📋 Meeting Assistant", elem_classes="main-header")
        gr.Markdown("Upload an audio or video file to transcribe, analyze, and extract meeting insights.")
        
        with gr.Tabs():
            with gr.TabItem("Process Meeting"):
                with gr.Row():
                    with gr.Column(scale=3):
                        meeting_date = gr.DateTime(
                            label="Meeting Date", 
                            value=datetime.now(), 
                        )
                        audio_in = gr.File(
                            type="filepath", 
                            label="Upload meeting audio/video",
                            file_types=SUPPORTED_FORMATS
                        )
                        gr.Markdown(
                            f"**Supported formats:** Audio: {', '.join(AUDIO_FORMATS)} | Video: {', '.join(VIDEO_FORMATS)}",
                            elem_classes="format-hint"
                        )
                        
                        file_info = gr.Markdown(elem_classes="file-info")
                        
                        with gr.Accordion("Advanced Settings", open=False):
                            with gr.Row():
                                whisper_dd = gr.Dropdown(
                                    ["base", "small", "medium", "large-v3"], 
                                    value=config["whisper_model"], 
                                    label="Whisper model"
                                )
                                speakers_in = gr.Number(
                                    label="# Speakers (0 = auto)", 
                                    value=0, 
                                    precision=0
                                )
                            with gr.Row():
                                llm_dd = gr.Dropdown(
                                    ["auto", "openai", "ollama"], 
                                    value=config["llm_backend"], 
                                    label="LLM backend"
                                )
                                language = gr.Text(
                                    label="Transcription language", 
                                    value=config["language"]
                                )
                            with gr.Row():
                                oa_model = gr.Text(
                                    label="OpenAI model", 
                                    value=config["openai_model"]
                                )
                                ol_model = gr.Text(
                                    label="Ollama model", 
                                    value=config["ollama_model"]
                                )
                            with gr.Row():
                                ol_url = gr.Text(
                                    label="Ollama URL", 
                                    value=config["ollama_url"]
                                )
                            export_cb = gr.CheckboxGroup(
                                ["json", "txt", "docx", "pdf"], 
                                value=config["export_formats"], 
                                label="Export formats"
                            )
                        
                        run_btn = gr.Button("Process Meeting", variant="primary")
                    
                    with gr.Column(scale=7):
                            
                        with gr.Row():
                            with gr.Column():
                                with gr.Tabs():
                                    with gr.TabItem("Summary"):
                                        summary_out = gr.Markdown(label="Meeting Summary")
                                    with gr.TabItem("Transcript"):
                                        transcript_out = gr.Markdown(label="Transcript")
                                    with gr.TabItem("Action Items"):
                                        actions_out = gr.Markdown(label="Action Items")
                                    with gr.TabItem("JSON"):
                                        json_out = gr.JSON(label="Result JSON")
                        
                        with gr.Row():
                            with gr.Column():
                                gr.Markdown("### Download Results")
                                with gr.Row():
                                    json_download = gr.File(label="JSON", visible=False)
                                    txt_download = gr.File(label="TXT", visible=False)
                                    docx_download = gr.File(label="DOCX", visible=False)
                                    pdf_download = gr.File(label="PDF", visible=False)
                                    ics_download = gr.File(label="Calendar", visible=False)
                            
                # Update file info when a file is uploaded
                audio_in.change(
                    fn=get_file_info,
                    inputs=audio_in,
                    outputs=[file_info]
                )
                
                run_btn.click(
                    fn=process,
                    inputs=[audio_in, meeting_date, whisper_dd, speakers_in, llm_dd, oa_model, ol_model, ol_url, language, export_cb],
                    outputs=[json_out, json_download, txt_download, docx_download, pdf_download, ics_download],
                ).then(
                    _postprocess, 
                    inputs=[json_out, json_download, txt_download, docx_download, pdf_download, ics_download],
                    outputs=[summary_out, transcript_out, actions_out, json_out, json_download, txt_download, docx_download, pdf_download, ics_download]
                )
            
            with gr.TabItem("Configuration"):
                gr.Markdown("### Configure Meeting Pipeline")
                
                with gr.Row():
                    with gr.Column():
                        whisper_model = gr.Dropdown(
                            ["base", "small", "medium", "large-v3"], 
                            value=config["whisper_model"], 
                            label="Default Whisper model"
                        )
                        llm_backend = gr.Dropdown(
                            ["auto", "openai", "ollama"], 
                            value=config["llm_backend"], 
                            label="Default LLM backend"
                        )
                        openai_model = gr.Text(
                            label="Default OpenAI model", 
                            value=config["openai_model"]
                        )
                        ollama_model = gr.Text(
                            label="Default Ollama model", 
                            value=config["ollama_model"]
                        )
                        ollama_url = gr.Text(
                            label="Default Ollama URL", 
                            value=config["ollama_url"]
                        )
                        language = gr.Text(
                            label="Default language", 
                            value=config["language"]
                        )
                        default_due_days = gr.Number(
                            label="Default due days for tasks", 
                            value=config["default_due_days"], 
                            precision=0
                        )
                    
                    with gr.Column():
                        max_retries = gr.Number(
                            label="Max API retries", 
                            value=config["max_retries"], 
                            precision=0
                        )
                        retry_delay = gr.Number(
                            label="Retry delay (seconds)", 
                            value=config["retry_delay"]
                        )
                        parallel_workers = gr.Number(
                            label="Parallel workers", 
                            value=config["parallel_workers"], 
                            precision=0
                        )
                        cache_enabled = gr.Checkbox(
                            label="Enable caching", 
                            value=config["cache_enabled"]
                        )
                        
                        with gr.Row():
                            save_btn = gr.Button("Save Configuration", variant="primary")
                            reset_btn = gr.Button("Reset to Defaults")
                
                config_status = gr.Text(label="Status", interactive=False)
                config_json = gr.JSON(label="Current Configuration", value=config)
                
                save_btn.click(
                    fn=save_config_ui,
                    inputs=[whisper_model, llm_backend, openai_model, ollama_model, ollama_url, language, 
                           max_retries, retry_delay, parallel_workers, cache_enabled, default_due_days],
                    outputs=[config_status, config_json]
                )
                
                def reset_config():
                    save_config(DEFAULT_CONFIG)
                    return f"Configuration reset to defaults", DEFAULT_CONFIG
                
                reset_btn.click(
                    fn=reset_config,
                    outputs=[config_status, config_json]
                )
            
            with gr.TabItem("Help"):
                gr.Markdown("""
                ### How to Use Meeting Assistant
                
                1. **Upload a File**: Click on the audio/video upload area and select a file from your computer.
                   - Supported audio formats: MP3, WAV, FLAC, M4A, AAC, OGG, WMA
                   - Supported video formats: MP4, MOV, AVI, MKV, WebM, FLV, WMV
                
                2. **Configure Settings** (optional):
                   - **Whisper Model**: Choose the size of the speech recognition model (larger is more accurate but slower)
                   - **Number of Speakers**: Specify how many speakers are in the meeting (0 for auto-detection)
                   - **LLM Backend**: Choose which AI service to use for analysis
                   - **Language**: Set the language for transcription (default: English)
                   - **Export Formats**: Select which formats to save the results in
                
                3. **Process the Meeting**: Click the "Process Meeting" button to start analysis.
                
                4. **View Results**: Once processing is complete, you can view:
                   - **Summary**: Key points and decisions from the meeting
                   - **Transcript**: Full text of the meeting with speaker labels
                   - **Action Items**: Tasks identified from the meeting with owners and due dates
                   - **JSON**: Raw data in JSON format
                
                5. **Download Results**: Download the meeting summary in your selected formats.
                
                ### Tips for Best Results
                
                - For better speaker diarization, specify the number of speakers if known
                - Use a larger Whisper model for better accuracy with longer meetings
                - Ensure the audio quality is good (minimal background noise, clear speech)
                - For video files, the audio track will be extracted automatically
                
                ### Troubleshooting
                
                - If processing fails, check that FFmpeg is installed on your system
                - For OpenAI backend, ensure you have a valid API key
                - For Ollama backend, ensure the Ollama service is running and the URL is correct
                """)
    
    demo.launch(share=True)
# ────────────────────────────────────────────────────────────────────── #
if __name__ == "__main__":
    _launch_ui()