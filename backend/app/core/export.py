import json
import logging
from pathlib import Path
from typing import Dict, Any, Optional
from datetime import datetime

logger = logging.getLogger(__name__)

# Optional imports for export formats
try:
    from docx import Document
except ImportError:
    Document = None

try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
except ImportError:
    canvas = SimpleDocTemplate = Paragraph = Spacer = getSampleStyleSheet = None

def export_to_json(data: Dict[str, Any], filename: str) -> Path:
    """Export meeting results to a JSON file."""
    path = Path(filename)
    path.parent.mkdir(parents=True, exist_ok=True)
    
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False, default=str)
    
    logger.info(f"Exported meeting data to JSON: {path}")
    return path

def export_to_txt(data: Dict[str, Any], filename: str) -> Path:
    """Export meeting results to a text file."""
    path = Path(filename)
    path.parent.mkdir(parents=True, exist_ok=True)
    
    with open(path, "w", encoding="utf-8") as f:
        f.write("╔═══════════════════════════════════════════════════════════════════╗\n")
        f.write("║                      MEETING SUMMARY REPORT                       ║\n")
        f.write("╚═══════════════════════════════════════════════════════════════════╝\n\n")
        
        # Meeting metadata
        f.write("=" * 70 + "\n")
        f.write("MEETING INFORMATION\n")
        f.write("=" * 70 + "\n\n")
        
        if "filename" in data:
            f.write(f"Meeting Name: {data['filename']}\n")
        
        if "created_at" in data:
            created_at = data["created_at"]
            if isinstance(created_at, str):
                f.write(f"Date: {created_at}\n")
            elif isinstance(created_at, datetime):
                f.write(f"Date: {created_at.strftime('%Y-%m-%d %H:%M:%S')}\n")
        
        if "status" in data:
            f.write(f"Status: {data['status'].upper()}\n")
        
        if "folder" in data:
            f.write(f"Folder: {data['folder']}\n")
        
        if "tags" in data:
            f.write(f"Tags: {data['tags']}\n")
        
        if "model_info" in data:
            model_info = data["model_info"]
            f.write(f"\nModel Configuration:\n")
            f.write(f"  - Configuration: {model_info.get('name', 'N/A')}\n")
            f.write(f"  - Language: {model_info.get('transcription_language', 'N/A')}\n")
            f.write(f"  - Speakers: {model_info.get('number_of_speakers', 'auto')}\n")
        
        # Speakers
        if "speakers" in data and data["speakers"]:
            f.write(f"\nSpeakers: ")
            speaker_names = [s.get('name', 'Unknown') for s in data["speakers"]]
            f.write(", ".join(speaker_names))
            f.write("\n")
        
        f.write("\n")
        
        # Summary
        f.write("=" * 70 + "\n")
        f.write("MEETING SUMMARY\n")
        f.write("=" * 70 + "\n\n")
        if "summary" in data:
            if isinstance(data["summary"], list):
                for point in data["summary"]:
                    f.write(f"• {point}\n")
            else:
                f.write(f"{data['summary']}\n")
        f.write("\n")
        
        # Action Items
        f.write("=" * 70 + "\n")
        f.write("ACTION ITEMS\n")
        f.write("=" * 70 + "\n\n")
        if "action_items" in data and data["action_items"]:
            for idx, item in enumerate(data["action_items"], 1):
                if isinstance(item, dict):
                    f.write(f"{idx}. {item.get('task', 'Unknown task')}\n")
                    if "owner" in item and item["owner"]:
                        f.write(f"   Owner: {item['owner']}\n")
                    if "due_date" in item and item["due_date"]:
                        f.write(f"   Due Date: {item['due_date']}\n")
                    if "status" in item and item["status"]:
                        f.write(f"   Status: {item['status'].capitalize()}\n")
                    if "priority" in item and item["priority"]:
                        f.write(f"   Priority: {item['priority'].capitalize()}\n")
                    if "notes" in item and item["notes"]:
                        f.write(f"   Notes: {item['notes']}\n")
                else:
                    f.write(f"{idx}. {item}\n")
                f.write("\n")
        else:
            f.write("No action items recorded.\n\n")
        
        # Decisions
        if "decisions" in data and data["decisions"]:
            f.write("=" * 70 + "\n")
            f.write("DECISIONS\n")
            f.write("=" * 70 + "\n\n")
            for decision in data["decisions"]:
                f.write(f"• {decision}\n")
            f.write("\n")
        
        # Open Questions
        if "open_questions" in data and data["open_questions"]:
            f.write("=" * 70 + "\n")
            f.write("OPEN QUESTIONS\n")
            f.write("=" * 70 + "\n\n")
            for question in data["open_questions"]:
                f.write(f"• {question}\n")
            f.write("\n")
        
        # Keywords
        if "keywords" in data and data["keywords"]:
            f.write("=" * 70 + "\n")
            f.write("KEYWORDS\n")
            f.write("=" * 70 + "\n\n")
            f.write(", ".join(data["keywords"]))
            f.write("\n\n")
        
        # Full Transcript
        if "transcript" in data:
            f.write("=" * 70 + "\n")
            f.write("FULL TRANSCRIPT\n")
            f.write("=" * 70 + "\n\n")
            f.write(data["transcript"])
            f.write("\n\n")
        
        # Footer
        f.write("=" * 70 + "\n")
        f.write(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        f.write("=" * 70 + "\n")
    
    logger.info(f"Exported meeting data to TXT: {path}")
    return path

def export_to_docx(data: Dict[str, Any], filename: str) -> Optional[Path]:
    """Export meeting results to a DOCX file."""
    if Document is None:
        logger.warning("python-docx not installed. Cannot export to DOCX.")
        return None
    
    path = Path(filename)
    path.parent.mkdir(parents=True, exist_ok=True)
    
    document = Document()
    
    # Title
    title = document.add_heading('Meeting Summary Report', 0)
    
    # Meeting Information Section
    document.add_heading('Meeting Information', level=1)
    
    if "filename" in data:
        p = document.add_paragraph()
        p.add_run('Meeting Name: ').bold = True
        p.add_run(data['filename'])
    
    if "created_at" in data:
        created_at = data["created_at"]
        p = document.add_paragraph()
        p.add_run('Date: ').bold = True
        if isinstance(created_at, str):
            p.add_run(created_at)
        elif isinstance(created_at, datetime):
            p.add_run(created_at.strftime('%Y-%m-%d %H:%M:%S'))
    
    if "status" in data:
        p = document.add_paragraph()
        p.add_run('Status: ').bold = True
        p.add_run(data['status'].upper())
    
    if "folder" in data:
        p = document.add_paragraph()
        p.add_run('Folder: ').bold = True
        p.add_run(data['folder'])
    
    if "tags" in data:
        p = document.add_paragraph()
        p.add_run('Tags: ').bold = True
        p.add_run(data['tags'])
    
    # Model Configuration
    if "model_info" in data:
        model_info = data["model_info"]
        p = document.add_paragraph()
        p.add_run('Model Configuration: ').bold = True
        p.add_run(f"{model_info.get('name', 'N/A')}")
        p = document.add_paragraph(f"  Language: {model_info.get('transcription_language', 'N/A')}", style='List Bullet')
        p = document.add_paragraph(f"  Speakers: {model_info.get('number_of_speakers', 'auto')}", style='List Bullet')
    
    # Speakers
    if "speakers" in data and data["speakers"]:
        p = document.add_paragraph()
        p.add_run('Speakers: ').bold = True
        speaker_names = [s.get('name', 'Unknown') for s in data["speakers"]]
        p.add_run(", ".join(speaker_names))
    
    document.add_paragraph()  # Add spacing
    
    # Summary Section
    document.add_heading('Meeting Summary', level=1)
    if "summary" in data:
        if isinstance(data["summary"], list):
            for point in data["summary"]:
                document.add_paragraph(point, style='List Bullet')
        else:
            document.add_paragraph(data["summary"])
    else:
        document.add_paragraph("No summary available.")
    
    # Action Items Section
    document.add_heading('Action Items', level=1)
    if "action_items" in data and data["action_items"]:
        for idx, item in enumerate(data["action_items"], 1):
            if isinstance(item, dict):
                # Task
                p = document.add_paragraph(style='List Number')
                p.add_run(item.get('task', 'Unknown task')).bold = True
                
                # Details in sub-bullets
                if "owner" in item and item["owner"]:
                    detail_p = document.add_paragraph(f"Owner: {item['owner']}", style='List Bullet 2')
                if "due_date" in item and item["due_date"]:
                    detail_p = document.add_paragraph(f"Due Date: {item['due_date']}", style='List Bullet 2')
                if "status" in item and item["status"]:
                    detail_p = document.add_paragraph(f"Status: {item['status'].capitalize()}", style='List Bullet 2')
                if "priority" in item and item["priority"]:
                    detail_p = document.add_paragraph(f"Priority: {item['priority'].capitalize()}", style='List Bullet 2')
                if "notes" in item and item["notes"]:
                    detail_p = document.add_paragraph(f"Notes: {item['notes']}", style='List Bullet 2')
            else:
                document.add_paragraph(str(item), style='List Number')
    else:
        document.add_paragraph("No action items recorded.")
    
    # Decisions Section
    if "decisions" in data and data["decisions"]:
        document.add_heading('Decisions', level=1)
        for decision in data["decisions"]:
            document.add_paragraph(decision, style='List Bullet')
    
    # Open Questions Section
    if "open_questions" in data and data["open_questions"]:
        document.add_heading('Open Questions', level=1)
        for question in data["open_questions"]:
            document.add_paragraph(question, style='List Bullet')
    
    # Keywords Section
    if "keywords" in data and data["keywords"]:
        document.add_heading('Keywords', level=1)
        document.add_paragraph(", ".join(data["keywords"]))
    
    # Full Transcript Section
    if "transcript" in data:
        document.add_page_break()
        document.add_heading('Full Transcript', level=1)
        document.add_paragraph(data["transcript"])
    
    # Footer
    document.add_paragraph()
    footer_p = document.add_paragraph()
    footer_p.add_run(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").italic = True
    
    document.save(str(path))
    logger.info(f"Exported meeting data to DOCX: {path}")
    return path

def export_to_pdf(data: Dict[str, Any], filename: str) -> Optional[Path]:
    """Export meeting results to a PDF file."""
    if SimpleDocTemplate is None:
        logger.warning("reportlab not installed. Cannot export to PDF.")
        return None
    
    path = Path(filename)
    path.parent.mkdir(parents=True, exist_ok=True)
    
    doc = SimpleDocTemplate(str(path), pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    # Title
    story.append(Paragraph("Meeting Summary Report", styles["Title"]))
    story.append(Spacer(1, 20))
    
    # Meeting Information Section
    story.append(Paragraph("Meeting Information", styles["Heading1"]))
    story.append(Spacer(1, 6))
    
    if "filename" in data:
        story.append(Paragraph(f"<b>Meeting Name:</b> {data['filename']}", styles["Normal"]))
    
    if "created_at" in data:
        created_at = data["created_at"]
        if isinstance(created_at, str):
            story.append(Paragraph(f"<b>Date:</b> {created_at}", styles["Normal"]))
        elif isinstance(created_at, datetime):
            story.append(Paragraph(f"<b>Date:</b> {created_at.strftime('%Y-%m-%d %H:%M:%S')}", styles["Normal"]))
    
    if "status" in data:
        story.append(Paragraph(f"<b>Status:</b> {data['status'].upper()}", styles["Normal"]))
    
    if "folder" in data:
        story.append(Paragraph(f"<b>Folder:</b> {data['folder']}", styles["Normal"]))
    
    if "tags" in data:
        story.append(Paragraph(f"<b>Tags:</b> {data['tags']}", styles["Normal"]))
    
    # Model Configuration
    if "model_info" in data:
        model_info = data["model_info"]
        story.append(Paragraph(f"<b>Model Configuration:</b> {model_info.get('name', 'N/A')}", styles["Normal"]))
        story.append(Paragraph(f"  • Language: {model_info.get('transcription_language', 'N/A')}", styles["Normal"]))
        story.append(Paragraph(f"  • Speakers: {model_info.get('number_of_speakers', 'auto')}", styles["Normal"]))
    
    # Speakers
    if "speakers" in data and data["speakers"]:
        speaker_names = [s.get('name', 'Unknown') for s in data["speakers"]]
        story.append(Paragraph(f"<b>Speakers:</b> {', '.join(speaker_names)}", styles["Normal"]))
    
    story.append(Spacer(1, 20))
    
    # Summary Section
    story.append(Paragraph("Meeting Summary", styles["Heading1"]))
    story.append(Spacer(1, 6))
    if "summary" in data:
        if isinstance(data["summary"], list):
            for point in data["summary"]:
                story.append(Paragraph(f"• {point}", styles["Normal"]))
        else:
            story.append(Paragraph(data["summary"], styles["Normal"]))
    else:
        story.append(Paragraph("No summary available.", styles["Normal"]))
    
    story.append(Spacer(1, 20))
    
    # Action Items Section
    story.append(Paragraph("Action Items", styles["Heading1"]))
    story.append(Spacer(1, 6))
    if "action_items" in data and data["action_items"]:
        for idx, item in enumerate(data["action_items"], 1):
            if isinstance(item, dict):
                task = item.get('task', 'Unknown task')
                story.append(Paragraph(f"<b>{idx}. {task}</b>", styles["Normal"]))
                
                if "owner" in item and item["owner"]:
                    story.append(Paragraph(f"   Owner: {item['owner']}", styles["Normal"]))
                if "due_date" in item and item["due_date"]:
                    story.append(Paragraph(f"   Due Date: {item['due_date']}", styles["Normal"]))
                if "status" in item and item["status"]:
                    story.append(Paragraph(f"   Status: {item['status'].capitalize()}", styles["Normal"]))
                if "priority" in item and item["priority"]:
                    story.append(Paragraph(f"   Priority: {item['priority'].capitalize()}", styles["Normal"]))
                if "notes" in item and item["notes"]:
                    story.append(Paragraph(f"   Notes: {item['notes']}", styles["Normal"]))
                story.append(Spacer(1, 8))
            else:
                story.append(Paragraph(f"{idx}. {item}", styles["Normal"]))
    else:
        story.append(Paragraph("No action items recorded.", styles["Normal"]))
    
    story.append(Spacer(1, 20))
    
    # Decisions Section
    if "decisions" in data and data["decisions"]:
        story.append(Paragraph("Decisions", styles["Heading1"]))
        story.append(Spacer(1, 6))
        for decision in data["decisions"]:
            story.append(Paragraph(f"• {decision}", styles["Normal"]))
        story.append(Spacer(1, 20))
    
    # Open Questions Section
    if "open_questions" in data and data["open_questions"]:
        story.append(Paragraph("Open Questions", styles["Heading1"]))
        story.append(Spacer(1, 6))
        for question in data["open_questions"]:
            story.append(Paragraph(f"• {question}", styles["Normal"]))
        story.append(Spacer(1, 20))
    
    # Keywords Section
    if "keywords" in data and data["keywords"]:
        story.append(Paragraph("Keywords", styles["Heading1"]))
        story.append(Spacer(1, 6))
        story.append(Paragraph(", ".join(data["keywords"]), styles["Normal"]))
        story.append(Spacer(1, 20))
    
    # Full Transcript Section (Condensed)
    if "transcript" in data:
        story.append(Paragraph("Full Transcript", styles["Heading1"]))
        story.append(Spacer(1, 6))
        # Truncate long transcripts for PDF
        transcript = data["transcript"]
        if len(transcript) > 10000:
            transcript = transcript[:10000] + "\n\n... (Transcript truncated for PDF export. Please use TXT or DOCX format for full transcript)"
        story.append(Paragraph(transcript.replace('\n', '<br/>'), styles["Normal"]))
    
    # Footer
    story.append(Spacer(1, 20))
    story.append(Paragraph(f"<i>Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</i>", styles["Normal"]))
    
    doc.build(story)
    logger.info(f"Exported meeting data to PDF: {path}")
    return path

def export_meeting_data(
    data: Dict[str, Any], 
    base_filename: str, 
    formats: list = ["json", "txt"]
) -> Dict[str, Optional[Path]]:
    """
    Export meeting data to multiple formats.
    
    Args:
        data: Meeting data dictionary
        base_filename: Base filename without extension
        formats: List of formats to export ("json", "txt", "docx", "pdf")
    
    Returns:
        Dictionary mapping format names to export file paths
    """
    results = {}
    
    for fmt in formats:
        try:
            if fmt.lower() == "json":
                results["json"] = export_to_json(data, f"{base_filename}.json")
            elif fmt.lower() == "txt":
                results["txt"] = export_to_txt(data, f"{base_filename}.txt")
            elif fmt.lower() == "docx":
                results["docx"] = export_to_docx(data, f"{base_filename}.docx")
            elif fmt.lower() == "pdf":
                results["pdf"] = export_to_pdf(data, f"{base_filename}.pdf")
            else:
                logger.warning(f"Unknown export format: {fmt}")
                results[fmt] = None
        except Exception as e:
            logger.error(f"Failed to export to {fmt}: {e}")
            results[fmt] = None
    
    return results
