import json
import logging
from datetime import datetime
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)

# Optional imports for export formats
try:
    from docx import Document
except ImportError:
    Document = None

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.pdfgen import canvas
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
except ImportError:
    canvas = SimpleDocTemplate = Paragraph = Spacer = getSampleStyleSheet = None


def export_to_json(data: dict[str, Any], filename: str) -> Path:
    """Export meeting results to a JSON file."""
    path = Path(filename)
    path.parent.mkdir(parents=True, exist_ok=True)

    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False, default=str)

    logger.info(f"Exported meeting data to JSON: {path}")
    return path


def export_to_txt(data: dict[str, Any], filename: str) -> Path:
    """Export meeting results to a text file."""
    path = Path(filename)
    path.parent.mkdir(parents=True, exist_ok=True)

    with open(path, "w", encoding="utf-8") as f:
        f.write("╔═══════════════════════════════════════════════════════════════════╗\n")
        f.write("║                      MEETING SUMMARY REPORT                       ║\n")
        f.write("╚═══════════════════════════════════════════════════════════════════╝\n\n")

        # Meeting metadata
        f.write("=" * 70 + "\n")
        f.write("MEETING INFORMATION\n")
        f.write("=" * 70 + "\n\n")

        if "filename" in data:
            f.write(f"Meeting Name: {data['filename']}\n")

        if "created_at" in data:
            created_at = data["created_at"]
            if isinstance(created_at, str):
                f.write(f"Date: {created_at}\n")
            elif isinstance(created_at, datetime):
                f.write(f"Date: {created_at.strftime('%Y-%m-%d %H:%M:%S')}\n")

        if "status" in data:
            f.write(f"Status: {data['status'].upper()}\n")

        if "folder" in data:
            f.write(f"Folder: {data['folder']}\n")

        if "tags" in data:
            f.write(f"Tags: {data['tags']}\n")

        if "model_info" in data:
            model_info = data["model_info"]
            f.write("\nModel Configuration:\n")
            f.write(f"  - Configuration: {model_info.get('name', 'N/A')}\n")
            f.write(f"  - Language: {model_info.get('transcription_language', 'N/A')}\n")
            f.write(f"  - Speakers: {model_info.get('number_of_speakers', 'auto')}\n")

        # Speakers
        if "speakers" in data and data["speakers"]:
            f.write("\nSpeakers: ")
            speaker_names = [s.get("name", "Unknown") for s in data["speakers"]]
            f.write(", ".join(speaker_names))
            f.write("\n")

        f.write("\n")

        # Notes
        if "notes" in data and data["notes"]:
            f.write("=" * 70 + "\n")
            f.write("MEETING NOTES\n")
            f.write("=" * 70 + "\n\n")
            f.write(f"{data['notes']}\n\n")

        # Summary
        f.write("=" * 70 + "\n")
        f.write("MEETING SUMMARY\n")
        f.write("=" * 70 + "\n\n")
        if "summary" in data:
            if isinstance(data["summary"], list):
                for point in data["summary"]:
                    f.write(f"• {point}\n")
            else:
                f.write(f"{data['summary']}\n")
        f.write("\n")

        # Action Items
        f.write("=" * 70 + "\n")
        f.write("ACTION ITEMS\n")
        f.write("=" * 70 + "\n\n")
        if "action_items" in data and data["action_items"]:
            for idx, item in enumerate(data["action_items"], 1):
                if isinstance(item, dict):
                    f.write(f"{idx}. {item.get('task', 'Unknown task')}\n")
                    if "owner" in item and item["owner"]:
                        f.write(f"   Owner: {item['owner']}\n")
                    if "due_date" in item and item["due_date"]:
                        f.write(f"   Due Date: {item['due_date']}\n")
                    if "status" in item and item["status"]:
                        f.write(f"   Status: {item['status'].capitalize()}\n")
                    if "priority" in item and item["priority"]:
                        f.write(f"   Priority: {item['priority'].capitalize()}\n")
                    if "notes" in item and item["notes"]:
                        f.write(f"   Notes: {item['notes']}\n")
                else:
                    f.write(f"{idx}. {item}\n")
                f.write("\n")
        else:
            f.write("No action items recorded.\n\n")

        # Decisions
        if "decisions" in data and data["decisions"]:
            f.write("=" * 70 + "\n")
            f.write("DECISIONS\n")
            f.write("=" * 70 + "\n\n")
            for decision in data["decisions"]:
                f.write(f"• {decision}\n")
            f.write("\n")

        # Open Questions
        if "open_questions" in data and data["open_questions"]:
            f.write("=" * 70 + "\n")
            f.write("OPEN QUESTIONS\n")
            f.write("=" * 70 + "\n\n")
            for question in data["open_questions"]:
                f.write(f"• {question}\n")
            f.write("\n")

        # Keywords
        if "keywords" in data and data["keywords"]:
            f.write("=" * 70 + "\n")
            f.write("KEYWORDS\n")
            f.write("=" * 70 + "\n\n")
            f.write(", ".join(data["keywords"]))
            f.write("\n\n")

        # Full Transcript
        if "transcript" in data:
            f.write("=" * 70 + "\n")
            f.write("FULL TRANSCRIPT\n")
            f.write("=" * 70 + "\n\n")
            f.write(data["transcript"])
            f.write("\n\n")

        # Footer
        f.write("=" * 70 + "\n")
        f.write(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        f.write("=" * 70 + "\n")

    logger.info(f"Exported meeting data to TXT: {path}")
    return path


def export_to_docx(data: dict[str, Any], filename: str) -> Path | None:
    """Export meeting results to a DOCX file."""
    if Document is None:
        logger.warning("python-docx not installed. Cannot export to DOCX.")
        return None

    path = Path(filename)
    path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()

    # Title
    title = document.add_heading("Meeting Summary Report", 0)

    # Meeting Information Section
    document.add_heading("Meeting Information", level=1)

    if "filename" in data:
        p = document.add_paragraph()
        p.add_run("Meeting Name: ").bold = True
        p.add_run(data["filename"])

    if "created_at" in data:
        created_at = data["created_at"]
        p = document.add_paragraph()
        p.add_run("Date: ").bold = True
        if isinstance(created_at, str):
            p.add_run(created_at)
        elif isinstance(created_at, datetime):
            p.add_run(created_at.strftime("%Y-%m-%d %H:%M:%S"))

    if "status" in data:
        p = document.add_paragraph()
        p.add_run("Status: ").bold = True
        p.add_run(data["status"].upper())

    if "folder" in data:
        p = document.add_paragraph()
        p.add_run("Folder: ").bold = True
        p.add_run(data["folder"])

    if "tags" in data:
        p = document.add_paragraph()
        p.add_run("Tags: ").bold = True
        p.add_run(data["tags"])

    # Model Configuration
    if "model_info" in data:
        model_info = data["model_info"]
        p = document.add_paragraph()
        p.add_run("Model Configuration: ").bold = True
        p.add_run(f"{model_info.get('name', 'N/A')}")
        p = document.add_paragraph(
            f"  Language: {model_info.get('transcription_language', 'N/A')}", style="List Bullet"
        )
        p = document.add_paragraph(f"  Speakers: {model_info.get('number_of_speakers', 'auto')}", style="List Bullet")

    # Speakers
    if "speakers" in data and data["speakers"]:
        p = document.add_paragraph()
        p.add_run("Speakers: ").bold = True
        speaker_names = [s.get("name", "Unknown") for s in data["speakers"]]
        p.add_run(", ".join(speaker_names))

    document.add_paragraph()  # Add spacing

    # Notes Section
    if "notes" in data and data["notes"]:
        document.add_heading("Meeting Notes", level=1)
        document.add_paragraph(data["notes"])
        document.add_paragraph()  # Add spacing

    # Summary Section
    document.add_heading("Meeting Summary", level=1)
    if "summary" in data:
        if isinstance(data["summary"], list):
            for point in data["summary"]:
                document.add_paragraph(point, style="List Bullet")
        else:
            document.add_paragraph(data["summary"])
    else:
        document.add_paragraph("No summary available.")

    # Action Items Section
    document.add_heading("Action Items", level=1)
    if "action_items" in data and data["action_items"]:
        for _idx, item in enumerate(data["action_items"], 1):
            if isinstance(item, dict):
                # Task
                p = document.add_paragraph(style="List Number")
                p.add_run(item.get("task", "Unknown task")).bold = True

                # Details in sub-bullets
                if "owner" in item and item["owner"]:
                    detail_p = document.add_paragraph(f"Owner: {item['owner']}", style="List Bullet 2")
                if "due_date" in item and item["due_date"]:
                    detail_p = document.add_paragraph(f"Due Date: {item['due_date']}", style="List Bullet 2")
                if "status" in item and item["status"]:
                    detail_p = document.add_paragraph(f"Status: {item['status'].capitalize()}", style="List Bullet 2")
                if "priority" in item and item["priority"]:
                    detail_p = document.add_paragraph(
                        f"Priority: {item['priority'].capitalize()}", style="List Bullet 2"
                    )
                if "notes" in item and item["notes"]:
                    detail_p = document.add_paragraph(f"Notes: {item['notes']}", style="List Bullet 2")
            else:
                document.add_paragraph(str(item), style="List Number")
    else:
        document.add_paragraph("No action items recorded.")

    # Decisions Section
    if "decisions" in data and data["decisions"]:
        document.add_heading("Decisions", level=1)
        for decision in data["decisions"]:
            document.add_paragraph(decision, style="List Bullet")

    # Open Questions Section
    if "open_questions" in data and data["open_questions"]:
        document.add_heading("Open Questions", level=1)
        for question in data["open_questions"]:
            document.add_paragraph(question, style="List Bullet")

    # Keywords Section
    if "keywords" in data and data["keywords"]:
        document.add_heading("Keywords", level=1)
        document.add_paragraph(", ".join(data["keywords"]))

    # Full Transcript Section
    if "transcript" in data:
        document.add_page_break()
        document.add_heading("Full Transcript", level=1)
        document.add_paragraph(data["transcript"])

    # Footer
    document.add_paragraph()
    footer_p = document.add_paragraph()
    footer_p.add_run(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").italic = True

    document.save(str(path))
    logger.info(f"Exported meeting data to DOCX: {path}")
    return path


def export_to_pdf(data: dict[str, Any], filename: str) -> Path | None:
    """Export meeting results to a PDF file."""
    if SimpleDocTemplate is None:
        logger.warning("reportlab not installed. Cannot export to PDF.")
        return None

    path = Path(filename)
    path.parent.mkdir(parents=True, exist_ok=True)

    doc = SimpleDocTemplate(str(path), pagesize=letter)
    styles = getSampleStyleSheet()
    story = []

    # Title
    story.append(Paragraph("Meeting Summary Report", styles["Title"]))
    story.append(Spacer(1, 20))

    # Meeting Information Section
    story.append(Paragraph("Meeting Information", styles["Heading1"]))
    story.append(Spacer(1, 6))

    if "filename" in data:
        story.append(Paragraph(f"<b>Meeting Name:</b> {data['filename']}", styles["Normal"]))

    if "created_at" in data:
        created_at = data["created_at"]
        if isinstance(created_at, str):
            story.append(Paragraph(f"<b>Date:</b> {created_at}", styles["Normal"]))
        elif isinstance(created_at, datetime):
            story.append(Paragraph(f"<b>Date:</b> {created_at.strftime('%Y-%m-%d %H:%M:%S')}", styles["Normal"]))

    if "status" in data:
        story.append(Paragraph(f"<b>Status:</b> {data['status'].upper()}", styles["Normal"]))

    if "folder" in data:
        story.append(Paragraph(f"<b>Folder:</b> {data['folder']}", styles["Normal"]))

    if "tags" in data:
        story.append(Paragraph(f"<b>Tags:</b> {data['tags']}", styles["Normal"]))

    # Model Configuration
    if "model_info" in data:
        model_info = data["model_info"]
        story.append(Paragraph(f"<b>Model Configuration:</b> {model_info.get('name', 'N/A')}", styles["Normal"]))
        story.append(Paragraph(f"  • Language: {model_info.get('transcription_language', 'N/A')}", styles["Normal"]))
        story.append(Paragraph(f"  • Speakers: {model_info.get('number_of_speakers', 'auto')}", styles["Normal"]))

    # Speakers
    if "speakers" in data and data["speakers"]:
        speaker_names = [s.get("name", "Unknown") for s in data["speakers"]]
        story.append(Paragraph(f"<b>Speakers:</b> {', '.join(speaker_names)}", styles["Normal"]))

    story.append(Spacer(1, 20))

    # Notes Section
    if "notes" in data and data["notes"]:
        story.append(Paragraph("Meeting Notes", styles["Heading1"]))
        story.append(Spacer(1, 6))
        story.append(Paragraph(data["notes"].replace("\n", "<br/>"), styles["Normal"]))
        story.append(Spacer(1, 20))

    # Summary Section
    story.append(Paragraph("Meeting Summary", styles["Heading1"]))
    story.append(Spacer(1, 6))
    if "summary" in data:
        if isinstance(data["summary"], list):
            for point in data["summary"]:
                story.append(Paragraph(f"• {point}", styles["Normal"]))
        else:
            story.append(Paragraph(data["summary"], styles["Normal"]))
    else:
        story.append(Paragraph("No summary available.", styles["Normal"]))

    story.append(Spacer(1, 20))

    # Action Items Section
    story.append(Paragraph("Action Items", styles["Heading1"]))
    story.append(Spacer(1, 6))
    if "action_items" in data and data["action_items"]:
        for idx, item in enumerate(data["action_items"], 1):
            if isinstance(item, dict):
                task = item.get("task", "Unknown task")
                story.append(Paragraph(f"<b>{idx}. {task}</b>", styles["Normal"]))

                if "owner" in item and item["owner"]:
                    story.append(Paragraph(f"   Owner: {item['owner']}", styles["Normal"]))
                if "due_date" in item and item["due_date"]:
                    story.append(Paragraph(f"   Due Date: {item['due_date']}", styles["Normal"]))
                if "status" in item and item["status"]:
                    story.append(Paragraph(f"   Status: {item['status'].capitalize()}", styles["Normal"]))
                if "priority" in item and item["priority"]:
                    story.append(Paragraph(f"   Priority: {item['priority'].capitalize()}", styles["Normal"]))
                if "notes" in item and item["notes"]:
                    story.append(Paragraph(f"   Notes: {item['notes']}", styles["Normal"]))
                story.append(Spacer(1, 8))
            else:
                story.append(Paragraph(f"{idx}. {item}", styles["Normal"]))
    else:
        story.append(Paragraph("No action items recorded.", styles["Normal"]))

    story.append(Spacer(1, 20))

    # Decisions Section
    if "decisions" in data and data["decisions"]:
        story.append(Paragraph("Decisions", styles["Heading1"]))
        story.append(Spacer(1, 6))
        for decision in data["decisions"]:
            story.append(Paragraph(f"• {decision}", styles["Normal"]))
        story.append(Spacer(1, 20))

    # Open Questions Section
    if "open_questions" in data and data["open_questions"]:
        story.append(Paragraph("Open Questions", styles["Heading1"]))
        story.append(Spacer(1, 6))
        for question in data["open_questions"]:
            story.append(Paragraph(f"• {question}", styles["Normal"]))
        story.append(Spacer(1, 20))

    # Keywords Section
    if "keywords" in data and data["keywords"]:
        story.append(Paragraph("Keywords", styles["Heading1"]))
        story.append(Spacer(1, 6))
        story.append(Paragraph(", ".join(data["keywords"]), styles["Normal"]))
        story.append(Spacer(1, 20))

    # Full Transcript Section (Condensed)
    if "transcript" in data:
        story.append(Paragraph("Full Transcript", styles["Heading1"]))
        story.append(Spacer(1, 6))
        # Truncate long transcripts for PDF
        transcript = data["transcript"]
        if len(transcript) > 10000:
            transcript = (
                transcript[:10000]
                + "\n\n... (Transcript truncated for PDF export. Please use TXT or DOCX format for full transcript)"
            )
        story.append(Paragraph(transcript.replace("\n", "<br/>"), styles["Normal"]))

    # Footer
    story.append(Spacer(1, 20))
    story.append(Paragraph(f"<i>Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</i>", styles["Normal"]))

    doc.build(story)
    logger.info(f"Exported meeting data to PDF: {path}")
    return path


def export_meeting_data(data: dict[str, Any], base_filename: str, formats: list = None) -> dict[str, Path | None]:
    """
    Export meeting data to multiple formats.

    Args:
        data: Meeting data dictionary
        base_filename: Base filename without extension
        formats: List of formats to export ("json", "txt", "docx", "pdf")

    Returns:
        Dictionary mapping format names to export file paths
    """
    if formats is None:
        formats = ["json", "txt"]
    results = {}

    for fmt in formats:
        try:
            if fmt.lower() == "json":
                results["json"] = export_to_json(data, f"{base_filename}.json")
            elif fmt.lower() == "txt":
                results["txt"] = export_to_txt(data, f"{base_filename}.txt")
            elif fmt.lower() == "docx":
                results["docx"] = export_to_docx(data, f"{base_filename}.docx")
            elif fmt.lower() == "pdf":
                results["pdf"] = export_to_pdf(data, f"{base_filename}.pdf")
            else:
                logger.warning(f"Unknown export format: {fmt}")
                results[fmt] = None
        except Exception as e:
            logger.error(f"Failed to export to {fmt}: {e}")
            results[fmt] = None

    return results


def export_project_to_json(data: dict[str, Any], filename: str) -> Path:
    """Export project data to a JSON file."""
    return export_to_json(data, filename)


def export_project_to_txt(data: dict[str, Any], filename: str) -> Path:
    """Export project data to a text file."""
    path = Path(filename)
    path.parent.mkdir(parents=True, exist_ok=True)

    project = data.get("project", {})
    metrics = data.get("metrics", {})
    meetings = data.get("meetings", [])
    milestones = data.get("milestones", [])
    action_items = data.get("action_items", [])
    members = data.get("members", [])
    notes = data.get("notes", [])

    def format_dt(value: Any) -> str:
        if value is None:
            return "N/A"
        if isinstance(value, datetime):
            return value.strftime("%Y-%m-%d %H:%M:%S")
        return str(value)

    with open(path, "w", encoding="utf-8") as f:
        f.write("╔═══════════════════════════════════════════════════════════════════╗\n")
        f.write("║                      PROJECT SUMMARY REPORT                       ║\n")
        f.write("╚═══════════════════════════════════════════════════════════════════╝\n\n")

        f.write("=" * 70 + "\n")
        f.write("PROJECT INFORMATION\n")
        f.write("=" * 70 + "\n\n")

        f.write(f"Name: {project.get('name', 'N/A')}\n")
        f.write(f"Status: {project.get('status', 'N/A')}\n")
        if project.get("folders"):
            f.write(f"Folders: {', '.join(project.get('folders', []))}\n")
        if project.get("description"):
            f.write(f"Description: {project.get('description')}\n")
        f.write(f"Start Date: {format_dt(project.get('start_date'))}\n")
        f.write(f"Target End Date: {format_dt(project.get('target_end_date'))}\n")
        f.write(f"Actual End Date: {format_dt(project.get('actual_end_date'))}\n")
        f.write(f"Created At: {format_dt(project.get('created_at'))}\n")
        f.write(f"Updated At: {format_dt(project.get('updated_at'))}\n\n")

        f.write("=" * 70 + "\n")
        f.write("KEY METRICS\n")
        f.write("=" * 70 + "\n\n")
        f.write(f"Meetings: {metrics.get('meeting_count', 0)}\n")
        f.write(
            f"Action Items: {metrics.get('completed_action_items', 0)}/{metrics.get('action_item_count', 0)} completed\n"
        )
        f.write(f"Members: {metrics.get('member_count', 0)}\n")
        f.write(f"Milestones: {metrics.get('milestone_count', 0)}\n\n")

        f.write("=" * 70 + "\n")
        f.write("MILESTONES\n")
        f.write("=" * 70 + "\n\n")
        if milestones:
            for idx, milestone in enumerate(milestones, 1):
                f.write(f"{idx}. {milestone.get('name', 'Untitled')}\n")
                f.write(f"   Status: {milestone.get('status', 'pending')}\n")
                f.write(f"   Due Date: {format_dt(milestone.get('due_date'))}\n")
                f.write(f"   Completed At: {format_dt(milestone.get('completed_at'))}\n")
                if milestone.get("description"):
                    f.write(f"   Description: {milestone.get('description')}\n")
                f.write("\n")
        else:
            f.write("No milestones recorded.\n\n")

        f.write("=" * 70 + "\n")
        f.write("TEAM MEMBERS\n")
        f.write("=" * 70 + "\n\n")
        if members:
            for idx, member in enumerate(members, 1):
                f.write(f"{idx}. {member.get('name', 'Unknown')}\n")
                if member.get("email"):
                    f.write(f"   Email: {member.get('email')}\n")
                if member.get("role"):
                    f.write(f"   Role: {member.get('role')}\n")
                f.write("\n")
        else:
            f.write("No members recorded.\n\n")

        f.write("=" * 70 + "\n")
        f.write("MEETINGS\n")
        f.write("=" * 70 + "\n\n")
        if meetings:
            for idx, meeting in enumerate(meetings, 1):
                f.write(f"{idx}. {meeting.get('title') or meeting.get('filename') or 'Untitled Meeting'}\n")
                f.write(f"   Date: {format_dt(meeting.get('meeting_date') or meeting.get('created_at'))}\n")
                f.write(f"   Status: {meeting.get('status', 'unknown')}\n")
                f.write(f"   Action Items: {meeting.get('action_items_count', 0)}\n")
                speakers = meeting.get("speakers") or []
                if speakers:
                    f.write(f"   Speakers: {', '.join(speakers)}\n")
                f.write("\n")
        else:
            f.write("No meetings recorded.\n\n")

        f.write("=" * 70 + "\n")
        f.write("ACTION ITEMS\n")
        f.write("=" * 70 + "\n\n")
        if action_items:
            for idx, item in enumerate(action_items, 1):
                f.write(f"{idx}. {item.get('task', 'Untitled Task')}\n")
                if item.get("owner"):
                    f.write(f"   Owner: {item.get('owner')}\n")
                if item.get("due_date"):
                    f.write(f"   Due Date: {item.get('due_date')}\n")
                if item.get("status"):
                    f.write(f"   Status: {item.get('status')}\n")
                if item.get("priority"):
                    f.write(f"   Priority: {item.get('priority')}\n")
                meeting_title = item.get("meeting_title") or item.get("meeting_filename")
                if meeting_title:
                    f.write(f"   Meeting: {meeting_title}\n")
                if item.get("notes"):
                    f.write(f"   Notes: {item.get('notes')}\n")
                f.write("\n")
        else:
            f.write("No action items recorded.\n\n")

        f.write("=" * 70 + "\n")
        f.write("PROJECT NOTES\n")
        f.write("=" * 70 + "\n\n")
        if notes:
            for idx, note in enumerate(notes, 1):
                f.write(f"{idx}. {note.get('title', 'Untitled Note')}\n")
                f.write(f"   Pinned: {'Yes' if note.get('pinned') else 'No'}\n")
                f.write(f"   Updated: {format_dt(note.get('updated_at'))}\n")
                if note.get("content"):
                    f.write(f"   Content:\n{note.get('content')}\n")
                attachments = note.get("attachments", [])
                if attachments:
                    f.write("   Attachments:\n")
                    for attachment in attachments:
                        f.write(f"     - {attachment.get('filename', 'attachment')}\n")
                f.write("\n")
        else:
            f.write("No project notes recorded.\n\n")

        f.write("=" * 70 + "\n")
        f.write(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        f.write("=" * 70 + "\n")

    logger.info(f"Exported project data to TXT: {path}")
    return path


def export_project_to_docx(data: dict[str, Any], filename: str) -> Path | None:
    """Export project data to a DOCX file."""
    if Document is None:
        logger.warning("python-docx not installed. Cannot export project to DOCX.")
        return None

    path = Path(filename)
    path.parent.mkdir(parents=True, exist_ok=True)

    project = data.get("project", {})
    metrics = data.get("metrics", {})
    meetings = data.get("meetings", [])
    milestones = data.get("milestones", [])
    action_items = data.get("action_items", [])
    members = data.get("members", [])
    notes = data.get("notes", [])

    def format_dt(value: Any) -> str:
        if value is None:
            return "N/A"
        if isinstance(value, datetime):
            return value.strftime("%Y-%m-%d %H:%M:%S")
        return str(value)

    document = Document()
    document.add_heading("Project Summary Report", 0)

    document.add_heading("Project Information", level=1)
    document.add_paragraph(f"Name: {project.get('name', 'N/A')}")
    document.add_paragraph(f"Status: {project.get('status', 'N/A')}")
    if project.get("folders"):
        document.add_paragraph(f"Folders: {', '.join(project.get('folders', []))}")
    if project.get("description"):
        document.add_paragraph(f"Description: {project.get('description')}")
    document.add_paragraph(f"Start Date: {format_dt(project.get('start_date'))}")
    document.add_paragraph(f"Target End Date: {format_dt(project.get('target_end_date'))}")
    document.add_paragraph(f"Actual End Date: {format_dt(project.get('actual_end_date'))}")
    document.add_paragraph(f"Created At: {format_dt(project.get('created_at'))}")
    document.add_paragraph(f"Updated At: {format_dt(project.get('updated_at'))}")

    document.add_heading("Key Metrics", level=1)
    document.add_paragraph(f"Meetings: {metrics.get('meeting_count', 0)}", style="List Bullet")
    document.add_paragraph(
        f"Action Items: {metrics.get('completed_action_items', 0)}/{metrics.get('action_item_count', 0)} completed",
        style="List Bullet",
    )
    document.add_paragraph(f"Members: {metrics.get('member_count', 0)}", style="List Bullet")
    document.add_paragraph(f"Milestones: {metrics.get('milestone_count', 0)}", style="List Bullet")

    document.add_heading("Milestones", level=1)
    if milestones:
        for milestone in milestones:
            milestone_title = milestone.get("name", "Untitled")
            document.add_paragraph(milestone_title, style="List Number")
            document.add_paragraph(f"Status: {milestone.get('status', 'pending')}", style="List Bullet 2")
            document.add_paragraph(f"Due Date: {format_dt(milestone.get('due_date'))}", style="List Bullet 2")
            document.add_paragraph(f"Completed At: {format_dt(milestone.get('completed_at'))}", style="List Bullet 2")
            if milestone.get("description"):
                document.add_paragraph(f"Description: {milestone.get('description')}", style="List Bullet 2")
    else:
        document.add_paragraph("No milestones recorded.")

    document.add_heading("Team Members", level=1)
    if members:
        for member in members:
            member_line = member.get("name", "Unknown")
            document.add_paragraph(member_line, style="List Bullet")
            if member.get("email"):
                document.add_paragraph(f"Email: {member.get('email')}", style="List Bullet 2")
            if member.get("role"):
                document.add_paragraph(f"Role: {member.get('role')}", style="List Bullet 2")
    else:
        document.add_paragraph("No members recorded.")

    document.add_heading("Meetings", level=1)
    if meetings:
        for meeting in meetings:
            title = meeting.get("title") or meeting.get("filename") or "Untitled Meeting"
            document.add_paragraph(title, style="List Number")
            document.add_paragraph(
                f"Date: {format_dt(meeting.get('meeting_date') or meeting.get('created_at'))}",
                style="List Bullet 2",
            )
            document.add_paragraph(f"Status: {meeting.get('status', 'unknown')}", style="List Bullet 2")
            document.add_paragraph(f"Action Items: {meeting.get('action_items_count', 0)}", style="List Bullet 2")
            speakers = meeting.get("speakers") or []
            if speakers:
                document.add_paragraph(f"Speakers: {', '.join(speakers)}", style="List Bullet 2")
    else:
        document.add_paragraph("No meetings recorded.")

    document.add_heading("Action Items", level=1)
    if action_items:
        for item in action_items:
            document.add_paragraph(item.get("task", "Untitled Task"), style="List Number")
            if item.get("owner"):
                document.add_paragraph(f"Owner: {item.get('owner')}", style="List Bullet 2")
            if item.get("due_date"):
                document.add_paragraph(f"Due Date: {item.get('due_date')}", style="List Bullet 2")
            if item.get("status"):
                document.add_paragraph(f"Status: {item.get('status')}", style="List Bullet 2")
            if item.get("priority"):
                document.add_paragraph(f"Priority: {item.get('priority')}", style="List Bullet 2")
            meeting_title = item.get("meeting_title") or item.get("meeting_filename")
            if meeting_title:
                document.add_paragraph(f"Meeting: {meeting_title}", style="List Bullet 2")
            if item.get("notes"):
                document.add_paragraph(f"Notes: {item.get('notes')}", style="List Bullet 2")
    else:
        document.add_paragraph("No action items recorded.")

    document.add_heading("Project Notes", level=1)
    if notes:
        for note in notes:
            document.add_paragraph(note.get("title", "Untitled Note"), style="List Number")
            document.add_paragraph(f"Pinned: {'Yes' if note.get('pinned') else 'No'}", style="List Bullet 2")
            document.add_paragraph(f"Updated: {format_dt(note.get('updated_at'))}", style="List Bullet 2")
            if note.get("content"):
                document.add_paragraph(note.get("content"))
            attachments = note.get("attachments", [])
            if attachments:
                document.add_paragraph("Attachments:", style="List Bullet 2")
                for attachment in attachments:
                    document.add_paragraph(attachment.get("filename", "attachment"), style="List Bullet 3")
    else:
        document.add_paragraph("No project notes recorded.")

    document.add_paragraph()
    document.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", style="Intense Quote")

    document.save(str(path))
    logger.info(f"Exported project data to DOCX: {path}")
    return path


def export_project_to_pdf(data: dict[str, Any], filename: str) -> Path | None:
    """Export project data to a PDF file."""
    if SimpleDocTemplate is None:
        logger.warning("reportlab not installed. Cannot export project to PDF.")
        return None

    path = Path(filename)
    path.parent.mkdir(parents=True, exist_ok=True)

    project = data.get("project", {})
    metrics = data.get("metrics", {})
    meetings = data.get("meetings", [])
    milestones = data.get("milestones", [])
    action_items = data.get("action_items", [])
    members = data.get("members", [])
    notes = data.get("notes", [])

    def format_dt(value: Any) -> str:
        if value is None:
            return "N/A"
        if isinstance(value, datetime):
            return value.strftime("%Y-%m-%d %H:%M:%S")
        return str(value)

    doc = SimpleDocTemplate(str(path), pagesize=letter)
    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph("Project Summary Report", styles["Title"]))
    story.append(Spacer(1, 20))

    story.append(Paragraph("Project Information", styles["Heading1"]))
    story.append(Spacer(1, 6))
    story.append(Paragraph(f"<b>Name:</b> {project.get('name', 'N/A')}", styles["Normal"]))
    story.append(Paragraph(f"<b>Status:</b> {project.get('status', 'N/A')}", styles["Normal"]))
    if project.get("folders"):
        story.append(Paragraph(f"<b>Folders:</b> {', '.join(project.get('folders', []))}", styles["Normal"]))
    if project.get("description"):
        story.append(Paragraph(f"<b>Description:</b> {project.get('description')}", styles["Normal"]))
    story.append(Paragraph(f"<b>Start Date:</b> {format_dt(project.get('start_date'))}", styles["Normal"]))
    story.append(Paragraph(f"<b>Target End Date:</b> {format_dt(project.get('target_end_date'))}", styles["Normal"]))
    story.append(Paragraph(f"<b>Actual End Date:</b> {format_dt(project.get('actual_end_date'))}", styles["Normal"]))
    story.append(Paragraph(f"<b>Created At:</b> {format_dt(project.get('created_at'))}", styles["Normal"]))
    story.append(Paragraph(f"<b>Updated At:</b> {format_dt(project.get('updated_at'))}", styles["Normal"]))
    story.append(Spacer(1, 20))

    story.append(Paragraph("Key Metrics", styles["Heading1"]))
    story.append(Spacer(1, 6))
    story.append(Paragraph(f"• Meetings: {metrics.get('meeting_count', 0)}", styles["Normal"]))
    story.append(
        Paragraph(
            f"• Action Items: {metrics.get('completed_action_items', 0)}/{metrics.get('action_item_count', 0)} completed",
            styles["Normal"],
        )
    )
    story.append(Paragraph(f"• Members: {metrics.get('member_count', 0)}", styles["Normal"]))
    story.append(Paragraph(f"• Milestones: {metrics.get('milestone_count', 0)}", styles["Normal"]))
    story.append(Spacer(1, 20))

    story.append(Paragraph("Milestones", styles["Heading1"]))
    story.append(Spacer(1, 6))
    if milestones:
        for milestone in milestones:
            story.append(Paragraph(f"• {milestone.get('name', 'Untitled')}", styles["Normal"]))
            story.append(Paragraph(f"  Status: {milestone.get('status', 'pending')}", styles["Normal"]))
            story.append(Paragraph(f"  Due Date: {format_dt(milestone.get('due_date'))}", styles["Normal"]))
            story.append(Paragraph(f"  Completed At: {format_dt(milestone.get('completed_at'))}", styles["Normal"]))
            if milestone.get("description"):
                story.append(Paragraph(f"  Description: {milestone.get('description')}", styles["Normal"]))
            story.append(Spacer(1, 8))
    else:
        story.append(Paragraph("No milestones recorded.", styles["Normal"]))
    story.append(Spacer(1, 20))

    story.append(Paragraph("Team Members", styles["Heading1"]))
    story.append(Spacer(1, 6))
    if members:
        for member in members:
            story.append(Paragraph(f"• {member.get('name', 'Unknown')}", styles["Normal"]))
            if member.get("email"):
                story.append(Paragraph(f"  Email: {member.get('email')}", styles["Normal"]))
            if member.get("role"):
                story.append(Paragraph(f"  Role: {member.get('role')}", styles["Normal"]))
            story.append(Spacer(1, 6))
    else:
        story.append(Paragraph("No members recorded.", styles["Normal"]))
    story.append(Spacer(1, 20))

    story.append(Paragraph("Meetings", styles["Heading1"]))
    story.append(Spacer(1, 6))
    if meetings:
        for meeting in meetings:
            title = meeting.get("title") or meeting.get("filename") or "Untitled Meeting"
            story.append(Paragraph(f"• {title}", styles["Normal"]))
            story.append(
                Paragraph(
                    f"  Date: {format_dt(meeting.get('meeting_date') or meeting.get('created_at'))}",
                    styles["Normal"],
                )
            )
            story.append(Paragraph(f"  Status: {meeting.get('status', 'unknown')}", styles["Normal"]))
            story.append(Paragraph(f"  Action Items: {meeting.get('action_items_count', 0)}", styles["Normal"]))
            speakers = meeting.get("speakers") or []
            if speakers:
                story.append(Paragraph(f"  Speakers: {', '.join(speakers)}", styles["Normal"]))
            story.append(Spacer(1, 8))
    else:
        story.append(Paragraph("No meetings recorded.", styles["Normal"]))
    story.append(Spacer(1, 20))

    story.append(Paragraph("Action Items", styles["Heading1"]))
    story.append(Spacer(1, 6))
    if action_items:
        for item in action_items:
            story.append(Paragraph(f"• {item.get('task', 'Untitled Task')}", styles["Normal"]))
            if item.get("owner"):
                story.append(Paragraph(f"  Owner: {item.get('owner')}", styles["Normal"]))
            if item.get("due_date"):
                story.append(Paragraph(f"  Due Date: {item.get('due_date')}", styles["Normal"]))
            if item.get("status"):
                story.append(Paragraph(f"  Status: {item.get('status')}", styles["Normal"]))
            if item.get("priority"):
                story.append(Paragraph(f"  Priority: {item.get('priority')}", styles["Normal"]))
            meeting_title = item.get("meeting_title") or item.get("meeting_filename")
            if meeting_title:
                story.append(Paragraph(f"  Meeting: {meeting_title}", styles["Normal"]))
            if item.get("notes"):
                story.append(Paragraph(f"  Notes: {item.get('notes')}", styles["Normal"]))
            story.append(Spacer(1, 8))
    else:
        story.append(Paragraph("No action items recorded.", styles["Normal"]))
    story.append(Spacer(1, 20))

    story.append(Paragraph("Project Notes", styles["Heading1"]))
    story.append(Spacer(1, 6))
    if notes:
        for note in notes:
            story.append(Paragraph(f"• {note.get('title', 'Untitled Note')}", styles["Normal"]))
            story.append(Paragraph(f"  Pinned: {'Yes' if note.get('pinned') else 'No'}", styles["Normal"]))
            story.append(Paragraph(f"  Updated: {format_dt(note.get('updated_at'))}", styles["Normal"]))
            if note.get("content"):
                story.append(Paragraph(note.get("content").replace("\n", "<br/>"), styles["Normal"]))
            attachments = note.get("attachments", [])
            if attachments:
                story.append(Paragraph("  Attachments:", styles["Normal"]))
                for attachment in attachments:
                    story.append(Paragraph(f"   - {attachment.get('filename', 'attachment')}", styles["Normal"]))
            story.append(Spacer(1, 8))
    else:
        story.append(Paragraph("No project notes recorded.", styles["Normal"]))

    story.append(Spacer(1, 20))
    story.append(Paragraph(f"<i>Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</i>", styles["Normal"]))

    doc.build(story)
    logger.info(f"Exported project data to PDF: {path}")
    return path
